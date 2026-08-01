"""Erstellt die Gebrauchsanweisung als Word-Dokument."""
import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# --- Styles anpassen ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)

# Hilfsfunktionen
def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()  # Abstand

# === INHALT ===

doc.add_heading('Russisch Vokabeltrainer — Gebrauchsanweisung', level=0)

doc.add_heading('Willkommen', level=1)
doc.add_paragraph(
    'Der Russisch Vokabeltrainer ist eine vielseitige Lern-App direkt im Browser. '
    'Sie brauchen nichts zu installieren — einfach die Datei index.html im Browser öffnen und loslegen.'
)
doc.add_paragraph('Die App bietet zwei Hauptfunktionen:')
doc.add_paragraph('Karteikarten — Vokabeln lernen mit virtuellen Lernkarten, die man umdrehen kann', style='List Bullet')
doc.add_paragraph('Texte vorlesen — Russische Texte werden Satz für Satz laut vorgelesen', style='List Bullet')
doc.add_paragraph(
    'Beide Funktionen nutzen die Sprachausgabe Ihres Browsers, um Ihnen die korrekte russische '
    'Aussprache vorzusprechen. So trainieren Sie gleichzeitig Lesen, Verstehen und Hören.'
)

# --- 1 ---
doc.add_heading('1. So starten Sie die App', level=1)
doc.add_paragraph('Öffnen Sie die Datei index.html in einem modernen Browser (Chrome, Edge oder Firefox).', style='List Number')
doc.add_paragraph(
    'Auf dem Startbildschirm sehen Sie:'
    , style='List Number')
doc.add_paragraph('Zwei große Schaltflächen: Karteikarten (das Kartensymbol) und Texte vorlesen (das Buchsymbol)', style='List Bullet 2')
doc.add_paragraph('Darunter die Stimmeinstellungen: Wählen Sie bereits hier Geschlecht, Stimme und Tempo für die Sprachausgabe — diese Einstellungen gelten für beide Modi.', style='List Bullet 2')
doc.add_paragraph('Oben rechts können Sie die Sprache der Benutzeroberfläche umschalten: DE (Deutsch), EN (Englisch) oder RU (Russisch).', style='List Number')

# --- 2 ---
doc.add_heading('2. Karteikarten — Vokabeln lernen', level=1)

doc.add_heading('2.1 Vokabeln laden', level=2)
doc.add_paragraph(
    'Klicken Sie auf Karteikarten. Sie haben nun zwei Möglichkeiten:'
)
p = doc.add_paragraph()
r = p.add_run('Neue Datei laden: ')
r.bold = True
p.add_run(
    'Laden Sie eine Textdatei mit Ihren Vokabeln. '
    'Sie können die Datei entweder per Drag & Drop in den Bereich ziehen oder durch Klick die Datei auswählen. '
    'Die App merkt sich den zuletzt verwendeten Ordner, sodass der Dateidialog beim nächsten Mal dort wieder öffnet (Chrome/Edge). '
    'Die App erkennt automatisch, welches Format Ihre Datei hat (siehe Abschnitt 3).'
)
p = doc.add_paragraph()
r = p.add_run('Gespeicherte Vokabeln laden: ')
r.bold = True
p.add_run(
    'Unterhalb des Datei-Upload-Bereichs sehen Sie die Liste Ihrer gespeicherten Vokabeln '
    'aus früheren Sitzungen (siehe Abschnitt 6). Sie können:'
)
doc.add_paragraph('Alle Vokabeln laden — Ihre gesamte Vokabelsammlung als Karteikarten', style='List Bullet')
doc.add_paragraph('Eine einzelne Lektion laden — klicken Sie auf den Titel (z.B. "Im Restaurant", "Antonyme"), '
                   'um nur die Vokabeln dieser Lektion zu lernen', style='List Bullet')

doc.add_heading('2.2 Einstellungen vor dem Start', level=2)
doc.add_paragraph('Nach dem Laden erscheint ein Einstellungsbildschirm:')
doc.add_paragraph('Lernmodus — Wählen Sie zwischen Frei, SM-2 oder FSRS (siehe Abschnitt 4)', style='List Bullet')
doc.add_paragraph('Auto-Vorlesen — Die Vorderseite wird automatisch vorgelesen', style='List Bullet')
doc.add_paragraph('Rückseite vorlesen — Auch die deutsche Übersetzung wird vorgelesen', style='List Bullet')
doc.add_paragraph('Klicken Sie auf Karteikarten starten, um zu beginnen.')

doc.add_heading('2.3 Während des Lernens', level=2)
doc.add_paragraph('Karte umdrehen: Klicken Sie auf die Karte oder drücken Sie die Leertaste.', style='List Bullet')
doc.add_paragraph('Nächste/Vorherige Karte: Benutzen Sie die Pfeil-Buttons oder die Pfeiltasten.', style='List Bullet')
doc.add_paragraph('Vorlesen: Klicken Sie auf den Lautsprecher-Button, um die aktuelle Karte vorlesen zu lassen.', style='List Bullet')
doc.add_paragraph('Bild hinzufügen: Klicken Sie auf den 📷-Button, um ein Bild von Ihrem Gerät zur aktuellen Karte hinzuzufügen. Das Bild erscheint auf der Vorderseite.', style='List Bullet')
doc.add_paragraph('Bild löschen: Fahren Sie mit der Maus über ein Bild auf der Karte — es erscheint ein kleiner ✕-Button, mit dem Sie das Bild wieder entfernen können.', style='List Bullet')
p = doc.add_paragraph(
    'Im Spaced-Repetition-Modus (SM-2 oder FSRS): Bewerten Sie nach dem Umdrehen, wie gut Sie die Karte wussten:'
)
doc.add_paragraph('Nochmal (Taste 1) — Wusste ich nicht', style='List Bullet 2')
doc.add_paragraph('Schwer (Taste 2) — Gerade so gewusst', style='List Bullet 2')
doc.add_paragraph('Gut (Taste 3) — Gewusst, aber musste nachdenken', style='List Bullet 2')
doc.add_paragraph('Leicht (Taste 4) — Sofort gewusst', style='List Bullet 2')

doc.add_heading('2.4 Konjugation und Deklination', level=2)
doc.add_paragraph(
    'Bei Verben und Nomen koennen auf der Karteikarte zusaetzliche Grammatik-Buttons erscheinen:'
)
p = doc.add_paragraph()
r = p.add_run('Konjugation ')
r.bold = True
p.add_run(
    '(bei Verben): Klicken Sie auf den Button, um die vollstaendige Konjugationstabelle zu sehen -- '
    'mit Praesens, Futur, Vergangenheit, Imperativ, Partizipien und Gerundium. '
    'Aspekt und Aspektpartner werden ebenfalls angezeigt.'
)
p = doc.add_paragraph()
r = p.add_run('Deklination ')
r.bold = True
p.add_run(
    '(bei Nomen): Klicken Sie auf den Button, um die Deklinationstabelle zu sehen -- '
    'mit allen sechs Kasus (Nominativ, Genitiv, Dativ, Akkusativ, Instrumental, Praepositiv) '
    'im Singular und Plural. Genus und Belebtheit werden angezeigt.'
)
doc.add_paragraph(
    'Die Tabellen werden im Vokabel-Editor (vokabel_editor.html) per KI erzeugt und automatisch '
    'mit den Karteikarten synchronisiert. Wenn eine konjugierte oder deklinierte Form auf der Karte steht, '
    'wird sie automatisch durch die Grundform (Infinitiv bzw. Nominativ Singular) ersetzt.'
)
doc.add_paragraph('Druecken Sie Escape oder klicken Sie auf das X, um die Tabelle zu schliessen.')

doc.add_heading('2.5 Die Kartenliste', level=2)
doc.add_paragraph(
    'Unterhalb der Karteikarte sehen Sie eine scrollbare Liste aller geladenen Karten. '
    'Klicken Sie auf eine Karte, um direkt dorthin zu springen. Der Lautsprecher-Button neben jeder Karte liest diese einzeln vor. '
    'Karten mit Bildern werden durch ein 🖼-Symbol in der Liste gekennzeichnet.'
)

# --- 3 ---
doc.add_heading('3. Das flexible Eingabeformat', level=1)
doc.add_paragraph(
    'Die App versteht drei Eingabeformate und erkennt automatisch, welches Sie verwenden. '
    'Sie können die Vokabellisten mit jedem einfachen Texteditor (Notepad, TextEdit usw.) erstellen und als .txt-Datei speichern.'
)

doc.add_heading('Format 1: Komma-getrennt (einfachstes Format)', level=2)
doc.add_paragraph('Das einfachste Format: Russisch und Deutsch durch ein Komma getrennt, eine Karte pro Zeile.')
add_code_block(doc, 'дом, Haus\nшкола, Schule\nкнига, Buch\nкошка, Katze\nсобака, Hund')
doc.add_paragraph('Vorderseite = Russisch (links vom Komma), Rückseite = Deutsch (rechts vom Komma).')

doc.add_heading('Format 2: Tab-getrennt (Quizlet-kompatibel)', level=2)
doc.add_paragraph(
    'Wenn Sie Vokabellisten von Quizlet exportieren, werden Vorder- und Rückseite durch einen Tabulator getrennt. '
    'Die App erkennt das automatisch.'
)
add_code_block(doc, 'дом\tHaus\nшкола\tSchule\nкнига\tBuch')
doc.add_paragraph('(Zwischen dem russischen und deutschen Wort steht ein Tabulator, also die Tab-Taste.)')

doc.add_heading('Format 3: Erweitertes Format mit Zusatzinformationen', level=2)

p = doc.add_paragraph()
r = p.add_run('Grammatikhinweise')
r.bold = True
doc.add_paragraph('Setzen Sie Grammatikregeln in runde Klammern am Anfang der Zeile:')
add_code_block(doc, '(в + Akkusativ - Richtung: wohin?)Я иду в школу., Ich gehe in die Schule.')
doc.add_paragraph('Der Grammatikhinweis wird als goldene Überschrift auf der Karte angezeigt.')

p = doc.add_paragraph()
r = p.add_run('Mehrzeilige Karten mit Beispielsätzen')
r.bold = True
doc.add_paragraph('Verwenden Sie \\n im Text, um Zeilenumbrüche auf der Karte zu erzeugen:')
add_code_block(doc, '(в + Akkusativ)\\n"Я иду в школу.", Ich gehe in die Schule.')

p = doc.add_paragraph()
r = p.add_run('Sprachausgabe steuern')
r.bold = True
doc.add_paragraph('Setzen Sie Text in Anführungszeichen, um festzulegen, was vorgelesen werden soll:')
add_code_block(doc, '1:00, "один час"')
doc.add_paragraph('Hier wird 1:00 angezeigt, aber die Sprachausgabe liest „один час" vor.')

p = doc.add_paragraph()
r = p.add_run('Bilder auf Karten')
r.bold = True
doc.add_paragraph('Fügen Sie Bilder hinzu mit dem Tag (img:Pfad):')
add_code_block(doc, '(img:bilder/dom.svg)дом, Haus')
doc.add_paragraph('Das Bild wird auf der Vorderseite über dem russischen Wort angezeigt. Unterstützt werden SVG, PNG, JPG usw.')

doc.add_heading('Texte zum Vorlesen', level=2)
doc.add_paragraph('Für den Modus Texte vorlesen verwenden Sie einfache Fließtexte. Absätze werden durch Leerzeilen getrennt:')
add_code_block(doc, 'Здравствуйте. Вот, пожалуйста, ваше меню.\n\nЯ буду минеральную воду без газа.\n\nМне, пожалуйста, чёрный чай с лимоном.')

# --- 4 ---
doc.add_heading('4. Lernmodi — Spaced Repetition', level=1)

doc.add_heading('Was ist Spaced Repetition?', level=2)
doc.add_paragraph(
    'Spaced Repetition ist eine wissenschaftlich erprobte Lernmethode: Karten, die Sie gut kennen, werden seltener wiederholt. '
    'Karten, die Ihnen schwerfallen, werden häufiger gezeigt. So nutzen Sie Ihre Lernzeit optimal.'
)

doc.add_heading('Die drei Modi im Überblick', level=2)
add_table(doc,
    ['Modus', 'Beschreibung', 'Empfohlen für'],
    [
        ['Frei', 'Selbst durch alle Karten blättern', 'Erstes Kennenlernen neuer Vokabeln'],
        ['SM-2', 'Klassisches System (seit den 1980ern)', 'Systematisches Langzeitlernen'],
        ['FSRS', 'Moderner, optimierter Algorithmus', 'Optimales Langzeitlernen'],
    ]
)

doc.add_heading('So funktioniert SM-2 / FSRS', level=2)
doc.add_paragraph('Die App zeigt Ihnen zuerst fällige Karten (Wiederholungen) und dann neue Karten (max. 20 pro Sitzung, einstellbar).', style='List Number')
doc.add_paragraph('Schauen Sie sich die Vorderseite an und versuchen Sie, die Antwort zu erinnern.', style='List Number')
doc.add_paragraph('Drehen Sie die Karte um (Klick oder Leertaste).', style='List Number')
doc.add_paragraph('Bewerten Sie sich ehrlich:', style='List Number')
doc.add_paragraph('Nochmal (Taste 1) — Wusste ich nicht. Die Karte erscheint bald wieder.', style='List Bullet 2')
doc.add_paragraph('Schwer (Taste 2) — Gerade so gewusst, musste lange ueberlegen.', style='List Bullet 2')
doc.add_paragraph('Gut (Taste 3) — Gewusst, aber musste kurz nachdenken. Das ist die empfohlene Standardbewertung bei richtiger Antwort.', style='List Bullet 2')
doc.add_paragraph('Leicht (Taste 4) — Sofort und muehelos gewusst, ohne jedes Zoegern.', style='List Bullet 2')
doc.add_paragraph('Die App berechnet automatisch, wann Sie jede Karte wieder sehen sollten.', style='List Number')

doc.add_heading('Die Bewertungs-Buttons verstehen', level=2)
doc.add_paragraph(
    'Unter jedem Button wird angezeigt, in wie vielen Tagen die Karte bei dieser Bewertung wieder erscheint. '
    'Diese Intervalle sind nicht fest — sie veraendern sich je nach Ihrer Lernhistorie mit der jeweiligen Karte.'
)

p = doc.add_paragraph()
r = p.add_run('Neue Karten: Kurze, eng beieinanderliegende Intervalle')
r.bold = True
doc.add_paragraph(
    'Wenn Sie eine Karte zum ersten Mal sehen, hat das System noch keine Informationen ueber Ihr Wissen. '
    'Deshalb sind die Intervalle anfangs kurz und liegen nah beieinander:'
)
add_table(doc,
    ['Button', 'Intervall bei einer neuen Karte'],
    [
        ['Nochmal', '1 Tag'],
        ['Schwer', '1 Tag'],
        ['Gut', '2 Tage'],
        ['Leicht', '6 Tage'],
    ]
)
doc.add_paragraph(
    'Das ist ganz normal und kein Fehler! Bei neuen Karten macht es noch keinen grossen Unterschied, '
    'ob Sie "Schwer" oder "Nochmal" druecken — die Karte kommt in jedem Fall sehr bald wieder.'
)

p = doc.add_paragraph()
r = p.add_run('Gelernte Karten: Wachsende, gut unterscheidbare Intervalle')
r.bold = True
doc.add_paragraph(
    'Je oefter Sie eine Karte erfolgreich wiederholen, desto mehr vergroessern und spreizen sich die Intervalle. '
    'Nach einigen Wiederholungen mit "Gut" koennte es so aussehen:'
)
add_table(doc,
    ['Button', 'Intervall nach mehreren Wiederholungen'],
    [
        ['Nochmal', '1 Tag'],
        ['Schwer', '3 Tage'],
        ['Gut', '9 Tage'],
        ['Leicht', '23 Tage'],
    ]
)
doc.add_paragraph(
    'Hier sehen Sie den grossen Unterschied: Eine gut gelernte Karte wird bei "Gut" erst in 9 Tagen '
    'wieder gezeigt, bei "Leicht" sogar erst in 23 Tagen. Bei "Nochmal" wird sie dagegen fast '
    'zurueckgesetzt und erscheint schon morgen wieder.'
)

p = doc.add_paragraph()
r = p.add_run('Wie wachsen die Intervalle?')
r.bold = True
doc.add_paragraph(
    'Das FSRS-System merkt sich fuer jede Karte, wie oft und wie gut Sie sie bewertet haben. '
    'Daraus berechnet es einen individuellen Stabilitaetswert:'
)
doc.add_paragraph('Jedes Mal, wenn Sie eine Karte mit "Gut" oder "Leicht" bewerten, steigt die Stabilitaet — das Intervall wird laenger.', style='List Bullet')
doc.add_paragraph('Bei "Schwer" waechst das Intervall langsamer.', style='List Bullet')
doc.add_paragraph('Bei "Nochmal" wird die Stabilitaet fast zurueckgesetzt — die Karte kommt wieder wie eine fast neue Karte.', style='List Bullet')

p = doc.add_paragraph()
r = p.add_run('Praktischer Tipp: ')
r.bold = True
p.add_run(
    'Bewerten Sie die meisten Karten mit "Gut" (Taste 3). Verwenden Sie "Leicht" nur bei Woertern, '
    'die Sie wirklich im Schlaf koennen. "Schwer" passt, wenn Sie unsicher waren, aber noch richtig lagen. '
    '"Nochmal" druecken Sie, wenn Ihnen die Antwort nicht eingefallen ist.'
)

doc.add_heading('Was bedeuten "Faellig", "Neu" und "Erledigt"?', level=2)
doc.add_paragraph('In der Statusleiste unten sehen Sie drei Zahlen:')
doc.add_paragraph('Faellig (rot) — Karten, die Sie schon einmal gelernt haben und die heute zur Wiederholung anstehen. Diese Karten haben Prioritaet.', style='List Bullet')
doc.add_paragraph('Neu (blau) — Karten, die Sie noch nie bewertet haben. Pro Sitzung werden maximal 20 neue Karten eingefuehrt.', style='List Bullet')
doc.add_paragraph('Erledigt (gruen/grau) — Karten, die Sie in dieser Sitzung bereits bearbeitet haben.', style='List Bullet')
doc.add_paragraph(
    'Wenn alle Karten gelernt sind und keine faellig sind, zeigt die App die Meldung '
    '"Keine Karten faellig und keine neuen Karten. Komm spaeter wieder!" '
    'Das bedeutet: Sie haben heute alles geschafft! Kommen Sie morgen (oder wann die naechsten Karten '
    'faellig werden) wieder, um weiterzulernen.'
)

doc.add_heading('Statistik und Fortschritt', level=2)
doc.add_paragraph(
    'Klicken Sie auf Statistik, um Ihren Lernfortschritt zu sehen: wie viele Karten Sie gelernt haben, '
    'wie viele Wiederholungen Sie durchgeführt haben, und Ihren täglichen Lernverlauf. '
    'Der Fortschritt wird automatisch im Browser gespeichert.'
)

# --- 5 ---
doc.add_heading('5. Texte vorlesen — Hörverstehen trainieren', level=1)
doc.add_paragraph('Klicken Sie auf Texte vorlesen und laden Sie eine Textdatei (Fließtext auf Russisch).', style='List Number')
doc.add_paragraph('Der Text wird in einem Lesefenster angezeigt.', style='List Number')
doc.add_paragraph('Wählen Sie den Lesemodus:', style='List Number')
doc.add_paragraph('ganzer Text — Der Text wird Satz für Satz vorgelesen. Das aktuell gesprochene Wort wird goldfarben markiert (Wort-für-Wort-Cursor). Bereits gelesene Sätze und Wörter werden ausgegraut.', style='List Bullet 2')
doc.add_paragraph('Satz für Satz — Jeder Satz wird einzeln vorgelesen; klicken Sie auf einen Satz, um ihn zu hören. Auch hier wird das aktuelle Wort markiert. Mit den Buttons ⏮ Voriger Satz / ⏭ Nächster Satz navigieren Sie zwischen den Sätzen.', style='List Bullet 2')
doc.add_paragraph(
    'Der Text scrollt automatisch mit, damit das aktuelle Wort immer im sichtbaren Bereich bleibt.'
)

# --- 6 ---
doc.add_heading('6. Die Vokabel-Datenbank — Ihr wachsender Wortschatz', level=1)

doc.add_heading('Automatisches Sammeln', level=2)
doc.add_paragraph(
    'Jedes Mal, wenn Sie eine Vokabeldatei laden oder Vokabeln per KI generieren lassen, werden die Karten automatisch in einer '
    'Datenbank im Browser gespeichert. Die Datenbank wächst mit jeder neuen Datei — ohne Duplikate. '
    'Jede Vokabel merkt sich, aus welcher Quelldatei (Lektion) sie stammt.'
)

doc.add_heading('Gespeicherte Vokabeln laden und lernen', level=2)
doc.add_paragraph(
    'Wenn Sie im Hauptmenü auf Karteikarten klicken, sehen Sie unterhalb des Datei-Upload-Bereichs Ihre gespeicherten Vokabeln:'
)
doc.add_paragraph('Alle Vokabeln (N Karten) — Lädt Ihre gesamte Vokabelsammlung als Karteikarten. Ideal für gemischtes Wiederholen.', style='List Bullet')
doc.add_paragraph('Einzelne Lektionen — Jede Datei, die Sie jemals geladen haben, erscheint als eigener Eintrag mit Kartenzähler. '
                   'Klicken Sie auf eine Lektion, um nur deren Vokabeln zu lernen.', style='List Bullet')
doc.add_paragraph(
    'So können Sie gezielt bestimmte Themen wiederholen, ohne die Originaldatei erneut laden zu müssen. '
    'Der Lernfortschritt (Spaced Repetition) bleibt dabei erhalten.'
)

doc.add_heading('Export für Excel', level=2)
doc.add_paragraph(
    'Klicken Sie im Einstellungsbildschirm auf Export (Excel), um Ihre gesamte Vokabelsammlung als CSV-Datei herunterzuladen. '
    'Diese Datei können Sie direkt in Excel oder Google Sheets öffnen.'
)
doc.add_paragraph('Die Exportdatei enthält: Russisch, Deutsch, Thema, Kategorie, Grammatik und Quelle (Dateiname).')

doc.add_heading('Auto-Export als JSON', level=2)
doc.add_paragraph(
    'In den Einstellungen (Zahnrad-Symbol oben rechts) können Sie einen Speicherort für eine JSON-Datei wählen. '
    'Die App exportiert dann bei jeder Änderung automatisch Ihre gesamte Vokabelsammlung in diese Datei. '
    'So haben Sie immer ein aktuelles Backup.'
)

# --- 7 ---
doc.add_heading('7. Tastenkürzel', level=1)
add_table(doc,
    ['Taste', 'Funktion'],
    [
        ['Leertaste', 'Karte umdrehen'],
        ['Pfeiltaste rechts', 'Nächste Karte'],
        ['Pfeiltaste links', 'Vorherige Karte'],
        ['1', 'Bewertung: Nochmal'],
        ['2', 'Bewertung: Schwer'],
        ['3', 'Bewertung: Gut'],
        ['4', 'Bewertung: Leicht'],
    ]
)

# --- 8 ---
doc.add_heading('8. KI und LLMs zur Erstellung von Vokabellisten', level=1)
doc.add_paragraph(
    'Eine der größten Stärken dieser App ist ihr einfaches Textformat. Sie können ChatGPT, Claude, Gemini oder andere '
    'KI-Assistenten bitten, Ihnen maßgeschneiderte Vokabellisten zu erstellen — in Sekunden, genau auf Ihr Niveau und Thema zugeschnitten.'
)

# 8.1
doc.add_heading('8.1 Einfache Vokabellisten erstellen', level=2)
p = doc.add_paragraph()
r = p.add_run('Prompt:')
r.bold = True
add_code_block(doc,
    'Erstelle mir eine Vokabelliste Russisch-Deutsch zum Thema\n'
    '"Essen und Trinken" für Anfänger (Niveau A1). Format: ein Eintrag\n'
    'pro Zeile, russisches Wort (mit Betonungszeichen), Komma,\n'
    'deutsches Wort.'
)
p = doc.add_paragraph()
r = p.add_run('Ergebnis (direkt nutzbar):')
r.bold = True
add_code_block(doc,
    'хлеб, Brot\nмолоко́, Milch\nмя́со, Fleisch\nры́ба, Fisch\n'
    'о́вощи, Gemüse\nфру́кты, Obst\nвода́, Wasser\nчай, Tee\n'
    'ко́фе, Kaffee\nсок, Saft\nсыр, Käse\nма́сло, Butter'
)

# 8.2
doc.add_heading('8.2 Vokabellisten mit Grammatikhinweisen', level=2)
p = doc.add_paragraph()
r = p.add_run('Prompt:')
r.bold = True
add_code_block(doc,
    'Erstelle mir 10 Karteikarten zum Thema "Russische Verben der\n'
    'Bewegung" mit Grammatikhinweisen. Format:\n\n'
    '(Grammatikhinweis)\\n"Russischer Beispielsatz", Deutsche Übersetzung\n\n'
    'Jede Karte soll einen Grammatikhinweis in Klammern am Anfang\n'
    'haben, dann \\n, dann den russischen Satz in Anführungszeichen\n'
    'und nach dem Komma die deutsche Übersetzung.'
)
p = doc.add_paragraph()
r = p.add_run('Ergebnis:')
r.bold = True
add_code_block(doc,
    '(идти́ - gehen, zu Fuß, einmalig)\\n"Я иду́ в магази́н.", Ich gehe in den Laden.\n'
    '(ходи́ть - gehen, regelmäßig)\\n"Я хожу́ в шко́лу ка́ждый день.", Ich gehe jeden Tag in die Schule.\n'
    '(е́хать - fahren, einmalig)\\n"Мы е́дем в Москву́.", Wir fahren nach Moskau.\n'
    '(е́здить - fahren, regelmäßig)\\n"Он е́здит на рабо́ту на метро́.", Er fährt mit der Metro zur Arbeit.\n'
    '(лете́ть - fliegen, einmalig)\\n"Самолёт лети́т в Берли́н.", Das Flugzeug fliegt nach Berlin.'
)

# 8.3
doc.add_heading('8.3 Vokabeln zu einem bestimmten Text erstellen', level=2)
p = doc.add_paragraph()
r = p.add_run('Prompt:')
r.bold = True
add_code_block(doc,
    'Hier ist ein russischer Text aus meinem Lehrbuch. Erstelle daraus\n'
    'eine Vokabelliste mit allen wichtigen Wörtern, die ein A1/A2-\n'
    'Lerner wahrscheinlich noch nicht kennt. Format: russisches Wort\n'
    'mit Betonungszeichen, Komma, deutsches Wort.\n\n'
    'Text: "Здравствуйте. Вот, пожалуйста, ваше меню. Какие напитки\n'
    'желаете?"'
)
p = doc.add_paragraph()
r = p.add_run('Ergebnis:')
r.bold = True
add_code_block(doc,
    'меню́, Speisekarte\nнапи́тки, Getränke\nжела́ть, wünschen\n'
    'минера́льная вода́, Mineralwasser\nбез газа́, ohne Kohlensäure\n'
    'чёрный чай, schwarzer Tee\nлимо́н, Zitrone'
)

# 8.4
doc.add_heading('8.4 Quizlet-Format (Tab-getrennt)', level=2)
p = doc.add_paragraph()
r = p.add_run('Prompt:')
r.bold = True
add_code_block(doc,
    'Erstelle eine Quizlet-kompatible Vokabelliste (Tab-getrennt) mit\n'
    '20 russischen Adjektiven und ihren deutschen Übersetzungen.\n'
    'Niveau A2. Format: Russisch [TAB] Deutsch.'
)
p = doc.add_paragraph()
r = p.add_run('Ergebnis (Tab-getrennt):')
r.bold = True
add_code_block(doc,
    'большо́й\tgroß\nма́ленький\tklein\nно́вый\tneu\nста́рый\talt\n'
    'молодо́й\tjung\nкраси́вый\tschön\nдо́брый\tgut / freundlich\n'
    'плохо́й\tschlecht\nбы́стрый\tschnell\nме́дленный\tlangsam'
)

# 8.5
doc.add_heading('8.5 Dialogkarten mit ganzen Sätzen', level=2)
p = doc.add_paragraph()
r = p.add_run('Prompt:')
r.bold = True
add_code_block(doc,
    'Erstelle 8 Karteikarten mit typischen Sätzen für eine Situation\n'
    'im Restaurant auf Russisch (Niveau A2). Jede Karte soll einen\n'
    'russischen Satz vorne und die deutsche Übersetzung hinten haben.\n'
    'Format: russisch, deutsch — ein Eintrag pro Zeile.'
)
p = doc.add_paragraph()
r = p.add_run('Ergebnis:')
r.bold = True
add_code_block(doc,
    'Мо́жно меню́, пожа́луйста?, Kann ich die Speisekarte haben, bitte?\n'
    'Что вы рекоменду́ете?, Was empfehlen Sie?\n'
    'Я бу́ду стейк с карто́фелем., Ich nehme das Steak mit Kartoffeln.\n'
    'Было́ о́чень вку́сно!, Es war sehr lecker!\n'
    'Счёт, пожа́луйста., Die Rechnung, bitte.\n'
    'Мо́жно плати́ть ка́ртой?, Kann ich mit Karte bezahlen?'
)

# 8.6
doc.add_heading('8.6 Thematische Karten mit Bildern', level=2)
p = doc.add_paragraph()
r = p.add_run('Prompt:')
r.bold = True
add_code_block(doc,
    'Erstelle eine Vokabelliste zum Thema "Familie" mit Bildern.\n'
    'Ich habe SVG-Bilder im Ordner bilder_familie/ mit folgenden\n'
    'Dateien: mama.svg, papa.svg, syn.svg, doch.svg.\n\n'
    'Format: (img:bilder_familie/datei.svg)russisch, deutsch'
)
p = doc.add_paragraph()
r = p.add_run('Ergebnis:')
r.bold = True
add_code_block(doc,
    '(img:bilder_familie/mama.svg)ма́ма, die Mama\n'
    '(img:bilder_familie/papa.svg)па́па, der Papa\n'
    '(img:bilder_familie/syn.svg)сын, der Sohn\n'
    '(img:bilder_familie/doch.svg)дочь, die Tochter'
)

# 8.7
doc.add_heading('8.7 Einen Vorlesetext erstellen lassen', level=2)
p = doc.add_paragraph()
r = p.add_run('Prompt:')
r.bold = True
add_code_block(doc,
    'Schreibe einen kurzen russischen Dialog (ca. 10 Sätze) zum Thema\n'
    '"Einkaufen im Supermarkt" für Niveau A2. Der Dialog soll zwischen\n'
    'einer Verkäuferin und einem Kunden stattfinden. Nur russischer\n'
    'Text, mit Absätzen zwischen den Sprecherwechseln.'
)
p = doc.add_paragraph()
r = p.add_run('Ergebnis:')
r.bold = True
add_code_block(doc,
    'Продавщица: Здравствуйте! Могу я вам помочь?\n\n'
    'Покупатель: Да, пожалуйста. Где у вас молочные продукты?\n\n'
    'Продавщица: Молоко и сыр в третьем ряду, справа.\n\n'
    'Покупатель: Спасибо. А свежий хлеб у вас есть?\n\n'
    'Продавщица: Да, конечно. Хлеб вот здесь, рядом с кассой.'
)
doc.add_paragraph(
    'Diesen Text speichern Sie als .txt-Datei und laden ihn im Modus Texte vorlesen — '
    'die App liest Ihnen den Dialog Satz für Satz vor, mit goldener Hervorhebung.'
)

# 8.8
doc.add_heading('8.8 Tipps für bessere KI-Ergebnisse', level=2)
doc.add_paragraph('Niveau angeben: Nennen Sie immer Ihr Sprachniveau (A1, A2, B1 usw.).', style='List Bullet')
doc.add_paragraph('Betonungszeichen verlangen: Schreiben Sie explizit „mit Betonungszeichen" in den Prompt.', style='List Bullet')
doc.add_paragraph('Format genau beschreiben: Geben Sie der KI ein Beispiel, wie eine Zeile aussehen soll.', style='List Bullet')
doc.add_paragraph('Thema eingrenzen: Je konkreter, desto nützlicher. Statt „Alltag" lieber „Tagesablauf eines Studenten".', style='List Bullet')
doc.add_paragraph('Direkt aus dem Lehrbuch: Kopieren Sie einen Text in den Prompt und bitten Sie die KI, Vokabeln zu extrahieren.', style='List Bullet')
doc.add_paragraph('Fehler prüfen: KI-generierte Inhalte gelegentlich mit einem Wörterbuch überprüfen.', style='List Bullet')

# --- 9 ---
doc.add_heading('9. Brauche ich einen API-Key?', level=1)
doc.add_paragraph(
    'Für den Vokabeltrainer (index.html) selbst: nein. Karteikarten lernen, Texte vorlesen, '
    'Spaced Repetition, Konjugations- und Deklinationstabellen anzeigen, Bilder, Export — '
    'alles funktioniert komplett offline ohne API-Key.'
)
doc.add_paragraph(
    'Ein Anthropic Claude API-Key wird ausschließlich im Vokabel-Editor (vokabel_editor.html) '
    'benötigt, und auch dort nur für die folgenden KI-Funktionen:'
)
add_table(doc,
    ['Funktion', 'Wozu?'],
    [
        ['Wortarten', 'Bestimmt automatisch, ob ein Wort Verb, Nomen, Adjektiv … ist'],
        ['Konjugation', 'Erzeugt vollständige Konjugationstabellen für Verben'],
        ['Deklination', 'Erzeugt Deklinationstabellen für Nomen (alle 6 Kasus, Sg./Pl.)'],
        ['Betonung', 'Markiert die betonte Silbe (wird vom Aussprachetrainer genutzt)'],
        ['Normalisieren', 'Bringt Vokabeln in Wörterbuchform'],
        ['AI-Vokabelfilter', 'Filtert beim Import von Texten Müll, Dubletten, Lateinschrift'],
    ]
)
doc.add_paragraph(
    'Alle KI-Funktionen erzeugen nur fehlende Daten — bereits vorhandene Konjugationen, '
    'Deklinationen usw. werden übersprungen, sodass der API-Verbrauch minimal bleibt.'
)
doc.add_paragraph(
    'Wenn Sie nur lernen und keine eigenen Vokabeln per KI anreichern möchten, brauchen Sie '
    'gar nichts einzurichten. Details zur API-Key-Konfiguration siehe Gebrauchsanweisung_Editor.md.'
)
doc.add_heading('Aussprachetrainer: Microsoft Azure', level=2)
doc.add_paragraph(
    'Der Aussprachetrainer (aussprache_trainer.html) nutzt zusätzlich den Azure Speech Service '
    'von Microsoft für die Bewertung Ihrer Aussprache auf Phonem-Ebene. Dafür brauchen Sie einen '
    'eigenen Azure-Key — nicht den Anthropic-Key.'
)
doc.add_paragraph('Anbieter: Microsoft Azure (kostenloses Free F0 Tier ist ausreichend)', style='List Bullet')
doc.add_paragraph('Benötigt: API-Key + Region (z.B. germanywestcentral)', style='List Bullet')
doc.add_paragraph('Eingabe einmalig im Setup-Bildschirm des Aussprachetrainers, danach im Browser gespeichert', style='List Bullet')
doc.add_paragraph(
    'Wird nur benötigt, wenn Sie Ihre Aussprache aufnehmen und bewerten lassen möchten. '
    'Das normale Vorlesen der Vokabeln (TTS im Vokabeltrainer) braucht keinen Azure-Key.',
    style='List Bullet'
)

doc.add_heading('Wo werden die API-Keys gespeichert?', level=2)
doc.add_paragraph(
    'Anthropic Claude-Key (Vokabeltrainer + Editor): Beide Apps suchen den Key in zwei Quellen, '
    'in dieser Reihenfolge:'
)
doc.add_paragraph(
    'Datei api-key.js im App-Ordner mit dem Inhalt: const CLAUDE_API_KEY = \'sk-ant-api03-IHR-ECHTER-KEY\'; '
    'Solange dort noch der Platzhalter YOUR_API_KEY_HERE steht, wird die Datei ignoriert. '
    'Sobald ein echter Key drinsteht, wird das Eingabefeld in den Einstellungen gesperrt und zeigt '
    '"********** (api-key.js)". Vorteil: Der Key liegt nicht im Browser-Speicher und ist über '
    '.gitignore von Versionskontrolle ausgeschlossen.',
    style='List Number'
)
doc.add_paragraph(
    'Browser-localStorage unter dem Schlüssel vokabeltrainer_apikey. Wird automatisch befüllt, '
    'sobald Sie den Key in das Eingabefeld der Einstellungen (index.html) bzw. der Filterleiste '
    '(vokabel_editor.html) eintragen.',
    style='List Number'
)
doc.add_paragraph(
    'Azure Speech-Key (Aussprachetrainer): Komplett getrennt vom Anthropic-Key und nicht in '
    'api-key.js. Wird beim ersten Setup im Aussprachetrainer eingegeben und im Browser-localStorage '
    'abgelegt:'
)
doc.add_paragraph('aussprache_azure_key — der API-Schlüssel', style='List Bullet')
doc.add_paragraph('aussprache_azure_region — die Azure-Region (z.B. germanywestcentral)', style='List Bullet')
doc.add_paragraph(
    'Keys einsehen oder löschen: Edge DevTools (F12) → Reiter Application → Storage → Local Storage → '
    'die Origin der App auswählen. Dort sind alle drei Einträge sichtbar und können gelöscht werden, '
    'falls Sie den Key wechseln oder von einem anderen Rechner aus neu setzen möchten.'
)

# --- 10 ---
doc.add_heading('10. Tipps und Hinweise', level=1)
doc.add_paragraph('Betonungszeichen (z.B. молоко́) helfen bei der Aussprache und werden korrekt angezeigt.', style='List Bullet')
doc.add_paragraph('Die verfügbaren russischen Stimmen hängen von Ihrem Browser ab. Chrome bietet die beste Auswahl.', style='List Bullet')
doc.add_paragraph('Alle Daten bleiben lokal in Ihrem Browser gespeichert. Es werden keine Daten an Server gesendet.', style='List Bullet')
doc.add_paragraph('Über die Modus-Wechsel-Buttons können Sie direkt zwischen Karteikarten und Text wechseln.', style='List Bullet')

# --- 11 ---
doc.add_heading('11. Zusammenfassung der Dateiformate', level=1)
add_table(doc,
    ['Format', 'Beispiel', 'Wann verwenden?'],
    [
        ['Komma-getrennt', 'дом, Haus', 'Einfache Vokabellisten'],
        ['Tab-getrennt', 'дом [TAB] Haus', 'Quizlet-Export, große Listen'],
        ['Mit Grammatik', '(Akk.)\\nSatz, Übersetzung', 'Grammatik + Beispielsätze'],
        ['Mit Bild', '(img:bild.svg)Wort, Übersetzung', 'Visuelles Lernen'],
        ['Fließtext', 'Absätze mit Leerzeilen', 'Texte vorlesen'],
    ]
)

doc.add_paragraph()

# ============================================================
# Anhang: Spaced-Repetition-Algorithmen
# ============================================================
doc.add_page_break()
doc.add_heading('Anhang: Die Spaced-Repetition-Algorithmen im Detail', level=1)

doc.add_paragraph(
    'Der Vokabeltrainer bietet zwei verschiedene Spaced-Repetition-Algorithmen: '
    'SM-2 (klassisch, seit den 1980er-Jahren in SuperMemo im Einsatz) und '
    'FSRS-4.5 (modern, datengetrieben, ab 2023). Beide haben dasselbe Ziel: '
    'Ihnen jede Karte genau dann zu zeigen, wenn Sie kurz davor sind, sie zu vergessen. '
    'Das nutzt den sogenannten Spacing-Effekt aus — Wiederholungen mit wachsenden '
    'Abständen festigen Wissen viel effizienter als stures Pauken.'
)
doc.add_paragraph(
    'Dieser Anhang erklärt, wie beide Algorithmen intern funktionieren. Sie müssen '
    'diese Details nicht verstehen, um die App zu benutzen — aber wenn Sie wissen '
    'wollen, warum eine Karte in 27 Tagen statt in 12 Tagen wiederkommt, finden Sie '
    'hier die Antwort.'
)

doc.add_heading('A.1 Grundidee von Spaced Repetition', level=2)
doc.add_paragraph(
    'Nach jeder Antwort speichert die App zu jeder Karte einen Datensatz mit folgenden Feldern:'
)
doc.add_paragraph('iterations / reps — wie oft die Karte schon bewertet wurde', style='List Bullet')
doc.add_paragraph('interval — Anzahl Tage bis zur nächsten Fälligkeit', style='List Bullet')
doc.add_paragraph('lastReview / nextReview — Zeitstempel (letzte Antwort / nächste Fälligkeit)', style='List Bullet')
doc.add_paragraph('easiness (nur SM-2) — individueller Schwierigkeits-Multiplikator der Karte', style='List Bullet')
doc.add_paragraph('difficulty / stability (nur FSRS) — zwei unabhängige Parameter (s. u.)', style='List Bullet')
doc.add_paragraph(
    'Die App zeigt Ihnen nur Karten, deren nextReview in der Vergangenheit liegt (fällig) '
    'oder die noch nie bewertet wurden (neu). Nach Ihrer Bewertung (Nochmal / Schwer / Gut / Leicht) '
    'wird interval neu berechnet und nextReview entsprechend gesetzt.'
)
doc.add_paragraph(
    'Die Rating-Buttons zeigen Ihnen schon vor dem Klick an, in wie vielen Tagen die Karte '
    'bei dieser Bewertung wieder käme — so können Sie die Konsequenzen abschätzen.'
)

doc.add_heading('A.2 SM-2 — Der Klassiker', level=2)
doc.add_paragraph(
    'SM-2 wurde 1985 von Piotr Woźniak für SuperMemo entwickelt und ist auch die Basis von Anki. '
    'Die Idee ist einfach: Jede Karte hat einen Easiness-Faktor (EF), der beschreibt, '
    'wie schwer sie Ihnen fällt. Beim Antworten wächst das Intervall multiplikativ mit EF.'
)

doc.add_heading('Die Intervall-Sequenz', level=3)
doc.add_paragraph('Bei einer korrekt beantworteten Karte (Gut oder Leicht):')
add_table(doc,
    ['Bewertung Nr.', 'Neues Intervall'],
    [
        ['1. (neu)', '1 Tag'],
        ['2.', '6 Tage'],
        ['3.', 'round(6 × EF) ≈ 16 Tage (bei EF=2.6)'],
        ['4.', 'round(vorheriges × EF)'],
        ['…', 'weiter multiplikativ'],
    ]
)
doc.add_paragraph(
    'Die ersten beiden Intervalle sind fest (1 und 6 Tage). Ab der dritten korrekten Antwort '
    'wird das Intervall mit dem aktuellen Easiness-Faktor der Karte multipliziert und gerundet.'
)

doc.add_heading('Der Easiness-Faktor', level=3)
doc.add_paragraph('Startwert: EF = 2.5 für jede neue Karte', style='List Bullet')
doc.add_paragraph('Untergrenze: EF darf niemals unter 1.3 fallen (sonst würden schwere Karten stagnieren)', style='List Bullet')
doc.add_paragraph("Aktualisierung nach korrekter Antwort: EF' = EF + (0.1 − (5 − q) · (0.08 + (5 − q) · 0.02))", style='List Bullet')
doc.add_paragraph('Dabei ist q die intern benutzte Qualitätsstufe:')
add_table(doc,
    ['Button', 'Qualität q', 'EF-Änderung'],
    [
        ['Nochmal', '1', 'EF wird nicht verändert; Karte wird zurückgesetzt'],
        ['Schwer', '3', '−0.14'],
        ['Gut', '4', '±0'],
        ['Leicht', '5', '+0.10'],
    ]
)
doc.add_paragraph(
    'Beispiel: Eine neue Karte startet mit EF = 2.5. Sie antworten dreimal hintereinander mit Leicht (q=5):'
)
doc.add_paragraph('Nach 1. Antwort: interval = 1, EF = 2.6', style='List Bullet')
doc.add_paragraph('Nach 2. Antwort: interval = 6, EF = 2.7', style='List Bullet')
doc.add_paragraph('Nach 3. Antwort: interval = round(6 × 2.7) = 16 Tage, EF = 2.8', style='List Bullet')
doc.add_paragraph('Nach 4. Antwort (16 Tage später): interval = round(16 × 2.8) = 45 Tage, EF = 2.9', style='List Bullet')

doc.add_heading('Fehlschlag (Nochmal)', level=3)
doc.add_paragraph(
    'Bei Nochmal wird die Karte komplett zurückgesetzt: iterations = 0, interval = 1. '
    'Sie kommt also am nächsten Tag wieder. Der Easiness-Faktor bleibt unverändert.'
)

doc.add_heading('Stärken und Schwächen', level=3)
doc.add_paragraph('Stärke: Sehr einfach, gut verstanden, seit Jahrzehnten bewährt', style='List Bullet')
doc.add_paragraph(
    'Schwäche: Die Formel kennt nur einen Parameter pro Karte. Sie unterscheidet nicht zwischen '
    '"schwer zu merken" und "wackelige Erinnerung" — beides fließt in den EF. FSRS macht das differenzierter.',
    style='List Bullet'
)

doc.add_heading('A.3 FSRS-4.5 — Der moderne Ansatz', level=2)
doc.add_paragraph(
    'FSRS (Free Spaced Repetition Scheduler) wurde ab 2022 von Jarrett Ye auf Basis der '
    'DSR-Theorie (Difficulty, Stability, Retrievability) entwickelt und trainiert seine '
    'Parameter an echten Review-Logs. Die Version 4.5 (Mitte 2023) ist inzwischen auch in '
    'Anki integriert und gilt als deutlich treffsicherer als SM-2.'
)

doc.add_heading('Die drei zentralen Größen', level=3)
doc.add_paragraph('FSRS beschreibt den Zustand einer Karte durch drei Werte:')
doc.add_paragraph(
    'Difficulty (D) — wie schwer die Karte dauerhaft für Sie ist (Skala 1–10, '
    'unabhängig davon wann Sie zuletzt geübt haben)',
    style='List Number'
)
doc.add_paragraph(
    'Stability (S) — wie lange die Erinnerung "hält"; konkret die Anzahl Tage, '
    'nach denen die Retrievability auf 90% fällt',
    style='List Number'
)
doc.add_paragraph(
    'Retrievability (R) — wie wahrscheinlich Sie die Karte jetzt gerade noch wissen '
    '(eine Zahl zwischen 0 und 1)',
    style='List Number'
)
doc.add_paragraph(
    'Der Clou: D und S sind unabhängig. Eine Karte kann schwer sein (hohes D), aber nach '
    'vielen Wiederholungen trotzdem extrem stabil werden (hohes S). SM-2 kann das nicht unterscheiden.'
)

doc.add_heading('Die Gewichte w[0..16]', level=3)
doc.add_paragraph(
    'FSRS-4.5 hat 17 globale Parameter (w[0] bis w[16]), die das Verhalten aller Karten steuern. '
    'Der Vokabeltrainer verwendet die Default-Gewichte des offiziellen FSRS-Projekts:'
)
add_code_block(doc,
    'w = [0.4, 0.6, 2.4, 5.8, 4.93, 0.94, 0.86, 0.01, 1.49,\n'
    '     0.14, 0.94, 2.18, 0.05, 0.34, 1.26, 0.29, 2.61]'
)
doc.add_paragraph('w[0..3] — Start-Stabilität nach der ersten Bewertung (Again/Hard/Good/Easy: 0.4/0.6/2.4/5.8 Tage)', style='List Bullet')
doc.add_paragraph('w[4], w[5] — Start-Schwierigkeit', style='List Bullet')
doc.add_paragraph('w[6] — wie stark jede Bewertung die Schwierigkeit verändert', style='List Bullet')
doc.add_paragraph('w[7] — Gewicht der Mean-Reversion (dazu gleich mehr)', style='List Bullet')
doc.add_paragraph('w[8..10] — Wachstum der Stabilität bei korrekter Antwort', style='List Bullet')
doc.add_paragraph('w[11..14] — Kollaps der Stabilität nach einem Vergessen', style='List Bullet')
doc.add_paragraph('w[15] — Penalty für Schwer (0.29, d. h. Stabilitätswachstum wird auf 29% gestaucht)', style='List Bullet')
doc.add_paragraph('w[16] — Bonus für Leicht (2.61, d. h. 2.6-fach beschleunigtes Wachstum)', style='List Bullet')

doc.add_heading('Initialisierung (erste Bewertung)', level=3)
doc.add_paragraph('Bei der allerersten Bewertung einer neuen Karte:')
doc.add_paragraph('Stabilität: S₀(G) = w[G−1], mit Untergrenze 0.1', style='List Bullet')
doc.add_paragraph('Again → 0.4 Tage, Hard → 0.6 Tage, Good → 2.4 Tage, Easy → 5.8 Tage', style='List Bullet 2')
doc.add_paragraph('Schwierigkeit: D₀(G) = clamp(w[4] − (G − 3) · w[5], 1, 10)', style='List Bullet')
doc.add_paragraph('Again → 6.81, Hard → 5.87, Good → 4.93, Easy → 3.99', style='List Bullet 2')
doc.add_paragraph('Das Intervall wird dann auf max(1, round(S)) gerundet.')

doc.add_heading('Retrievability — die Vergessenskurve', level=3)
doc.add_paragraph(
    'FSRS-4.5 modelliert das Vergessen nicht exponentiell (wie SM-2 implizit), sondern '
    'mit einer Potenzfunktion:'
)
add_code_block(doc, 'R(t, S) = (1 + t / (9·S))⁻¹')
doc.add_paragraph(
    'Dabei ist t die seit der letzten Antwort vergangene Zeit in Tagen und S die aktuelle '
    'Stabilität. Diese Formel hat zwei wichtige Eigenschaften:'
)
doc.add_paragraph('Bei t = 0: R = 1 (100% — Sie haben die Karte gerade gesehen)', style='List Bullet')
doc.add_paragraph('Bei t = S: R = 1 / (1 + 1/9) = 9/10 = 0.9', style='List Bullet')
doc.add_paragraph(
    'Der zweite Punkt ist die Definition der Stabilität: Stability ist die Zeit, nach der Sie '
    'die Karte noch mit 90% Wahrscheinlichkeit wissen. FSRS plant jede Wiederholung genau so, '
    'dass Sie kurz vor diesem 90%-Punkt stehen — das ist optimal für den Lerneffekt.'
)
doc.add_paragraph(
    'Die Potenzfunktion fällt langsamer ab als eine reine Exponentialfunktion. Das passt besser '
    'zu empirischen Gedächtnisdaten: Wenn Sie eine Karte einmal lange wissen, verfällt sie nicht '
    'abrupt, sondern gleitet flacher.'
)

doc.add_heading('Schwierigkeits-Update (bei jedem Review)', level=3)
doc.add_paragraph('Nach jeder Bewertung wird die Schwierigkeit in zwei Schritten angepasst:')
doc.add_paragraph('Schritt 1 — Linear Damping:')
add_code_block(doc, 'ΔD = −w[6] · (G − 3)\nD\' = D + ΔD · (10 − D) / 9')
doc.add_paragraph(
    'Der Faktor (10 − D) / 9 sorgt dafür, dass schon-sehr-schwere Karten (D nahe 10) nur '
    'noch schwach weiter ansteigen können — eine Karte darf nicht unbegrenzt in die Hölle wandern.'
)
doc.add_paragraph('Schritt 2 — Mean Reversion:')
add_code_block(doc, "D_neu = w[7] · D₀(4) + (1 − w[7]) · D'")
doc.add_paragraph(
    'Die Schwierigkeit wird ganz leicht (Gewicht w[7] = 0.01, also 1%) in Richtung der '
    'Start-Schwierigkeit bei Easy zurückgezogen. Das verhindert ein dauerhaftes "Driften" '
    'der Schwierigkeit. Anschließend wird D auf den Bereich [1, 10] geklemmt.'
)

doc.add_heading('Stabilitäts-Update bei korrekter Antwort (Hard/Good/Easy)', level=3)
doc.add_paragraph(
    'Das ist die Kernformel von FSRS. Sie berechnet, wie stark sich die Stabilität durch '
    'diese Wiederholung vergrößert:'
)
add_code_block(doc,
    'S_neu = S · (1 + exp(w[8]) · (11 − D) · S^(−w[9]) ·\n'
    '              (exp((1 − R) · w[10]) − 1) · hardPenalty · easyBonus)'
)
doc.add_paragraph('Was passiert da?')
doc.add_paragraph('exp(w[8]) — Basiswachstumsfaktor (≈ 4.4)', style='List Bullet')
doc.add_paragraph('(11 − D) — leichte Karten (kleines D) wachsen stärker als schwere', style='List Bullet')
doc.add_paragraph('S^(−w[9]) — schon-stabile Karten wachsen relativ langsamer (damit extrem stabile Karten nicht explodieren)', style='List Bullet')
doc.add_paragraph(
    '(exp((1 − R) · w[10]) − 1) — das ist die wichtigste Komponente: Je niedriger R beim Review '
    '(also je näher am Vergessen Sie waren), desto stärker wächst die Stabilität. Wer knapp davor '
    'war zu vergessen, lernt am meisten.',
    style='List Bullet'
)
doc.add_paragraph('hardPenalty = w[15] = 0.29 bei Schwer, sonst 1', style='List Bullet')
doc.add_paragraph('easyBonus = w[16] = 2.61 bei Leicht, sonst 1', style='List Bullet')
doc.add_paragraph(
    'Der berühmte Desirable-Difficulty-Effekt aus der Lernpsychologie steckt mathematisch in '
    'diesem (1 − R)-Term: Wiederholungen am Limit sind am wertvollsten.'
)

doc.add_heading('Stabilitäts-Kollaps bei Vergessen (Again)', level=3)
doc.add_paragraph(
    'Wenn Sie Nochmal drücken, bricht die Stabilität nicht auf null ein (wie bei SM-2), '
    'sondern kollabiert kontrolliert:'
)
add_code_block(doc, 'S_forget = w[11] · D^(−w[12]) · ((S + 1)^w[13] − 1) · exp((1 − R) · w[14])')
doc.add_paragraph('Schwere Karten verlieren mehr (D^(−w[12]))', style='List Bullet')
doc.add_paragraph('Aber hochstabile Karten bleiben relativ robust (der (S+1)^w[13]-Term)', style='List Bullet')
doc.add_paragraph('Und wenn Sie beim Vergessen noch eine hohe R hatten (also unerwartet vergessen), wird stärker bestraft', style='List Bullet')
doc.add_paragraph(
    'Das ist realistischer als SM-2: Eine Karte, die Sie seit zwei Jahren kennen und einmal '
    'vergessen, startet nicht wieder bei Tag 1.'
)

doc.add_heading('Intervallberechnung', level=3)
doc.add_paragraph('Nach jedem Update wird das neue Intervall gesetzt als:')
add_code_block(doc, 'interval = max(1, round(S_neu))')
doc.add_paragraph(
    'Also: Ein Tag pro Einheit Stabilität. Die Rating-Buttons zeigen Ihnen vorab die '
    'Vorhersage für alle vier Bewertungen, indem sie die Berechnung mit der aktuellen '
    'D/S/R einmal pro Button durchspielen.'
)

doc.add_heading('A.4 SM-2 oder FSRS — was soll ich wählen?', level=2)
doc.add_paragraph(
    'SM-2 ist berechenbar, transparent und funktioniert seit Jahrzehnten. Wenn Sie einfache '
    'Intervalle bevorzugen und keine Lust auf Black-Box-Formeln haben, ist SM-2 eine solide Wahl.',
    style='List Bullet'
)
doc.add_paragraph(
    'FSRS-4.5 passt sich besser an Ihren individuellen Lernverlauf an, unterscheidet zwischen '
    '"schwer" und "wackelig" und ist besonders stark bei Karten, die Sie schon lange kennen. '
    'Für Vokabeln mit Langzeit-Horizont ist FSRS in der Regel effizienter.',
    style='List Bullet'
)
doc.add_paragraph(
    'Die App speichert die Spaced-Repetition-Daten getrennt für beide Abfragerichtungen '
    '(Russisch→Deutsch und Deutsch→Russisch), sodass Sie jede Richtung unabhängig trainieren '
    'können — selbst mit unterschiedlichen Algorithmen.'
)

doc.add_heading('A.5 Wo stehen die Algorithmen im Code?', level=2)
doc.add_paragraph(
    'Beide Implementierungen liegen in js/sr.js. Die Datei ist bewusst klein gehalten '
    '(~100 Zeilen) und enthält keine DOM- oder State-Abhängigkeiten, sodass sie direkt '
    'getestet werden kann. Die zugehörigen Tests stehen in test_sr.js (44 Tests, '
    'Aufruf: node test_sr.js) und verifizieren die Intervall-Sequenzen, die '
    'Monotonie-Eigenschaften (Easy ≥ Good ≥ Hard > Again) und die korrekte Umsetzung '
    'der FSRS-4.5-Formeln.'
)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Viel Erfolg beim Russischlernen!')
r.bold = True
r.font.size = Pt(14)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Speichern
# M-26: Pfad relativ zum Skript statt fest verdrahtet -- das Projekt
# laesst sich damit verschieben, ohne die Skripte anzufassen.
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(BASE_DIR, '..', 'docs', 'Gebrauchsanweisung.docx')
try:
    doc.save(output_path)
except Exception as e:
    print('FEHLER: Word-Dokument konnte nicht gespeichert werden (%s): %s' % (output_path, e), file=sys.stderr)
    sys.exit(1)
print(f'Word-Dokument gespeichert: {output_path}')
