"""Creates the User Guide as a Word document (English version)."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()

# === CONTENT ===

doc.add_heading('Russian Vocabulary Trainer — User Guide', level=0)

doc.add_heading('Welcome', level=1)
doc.add_paragraph(
    'The Russian Vocabulary Trainer is a versatile learning app that runs directly in your browser. '
    'No installation required — just open the file index.html in your browser and start learning.'
)
doc.add_paragraph('The app offers two main features:')
doc.add_paragraph('Flashcards — Learn vocabulary with virtual cards you can flip', style='List Bullet')
doc.add_paragraph('Text Reader — Russian texts are read aloud sentence by sentence', style='List Bullet')
doc.add_paragraph(
    'Both features use your browser\'s speech synthesis to pronounce Russian correctly. '
    'This way you train reading, comprehension, and listening all at once.'
)

# --- 1 ---
doc.add_heading('1. Getting Started', level=1)
doc.add_paragraph('Open the file index.html in a modern browser (Chrome, Edge, or Firefox).', style='List Number')
doc.add_paragraph('On the start screen you will see two large buttons: Flashcards and Read Texts.', style='List Number')
doc.add_paragraph('At the top you can switch the interface language: DE (German), EN (English), or RU (Russian).', style='List Number')

# --- 2 ---
doc.add_heading('2. Flashcards — Learning Vocabulary', level=1)

doc.add_heading('2.1 Loading a File', level=2)
doc.add_paragraph(
    'Click on Flashcards and load a text file with your vocabulary. '
    'You can either drag & drop the file into the upload area or click to select it. '
    'The app remembers the last folder you used, so the file dialog opens there next time (Chrome/Edge). '
    'The app automatically detects which format your file uses (see Section 3).'
)

doc.add_heading('2.2 Settings Before Starting', level=2)
doc.add_paragraph('After loading, a settings screen appears:')
doc.add_paragraph('Learning mode — Choose between Free, SM-2, or FSRS (see Section 4)', style='List Bullet')
doc.add_paragraph('Voice — Select a male or female voice', style='List Bullet')
doc.add_paragraph('Speed — Adjust the speech rate (slower = easier to follow)', style='List Bullet')
doc.add_paragraph('Auto-read — The front side is read aloud automatically', style='List Bullet')
doc.add_paragraph('Read back side — The translation is also read aloud', style='List Bullet')
doc.add_paragraph('Click Start Flashcards to begin.')

doc.add_heading('2.3 During the Learning Session', level=2)
doc.add_paragraph('Flip the card: Click on the card or press the Spacebar.', style='List Bullet')
doc.add_paragraph('Next/Previous card: Use the arrow buttons or arrow keys.', style='List Bullet')
doc.add_paragraph('Read aloud: Click the speaker button to hear the current card.', style='List Bullet')
doc.add_paragraph('Add image: Click the 📷 button to add an image from your device to the current card. The image appears on the front side.', style='List Bullet')
doc.add_paragraph('Remove image: Hover over an image on the card — a small ✕ button appears to remove the image.', style='List Bullet')
doc.add_paragraph(
    'In Spaced Repetition mode (SM-2 or FSRS): After flipping, rate how well you knew the card:'
)
doc.add_paragraph('Again (Key 1) — Didn\'t know it', style='List Bullet 2')
doc.add_paragraph('Hard (Key 2) — Barely remembered', style='List Bullet 2')
doc.add_paragraph('Good (Key 3) — Knew it, but had to think', style='List Bullet 2')
doc.add_paragraph('Easy (Key 4) — Knew it instantly', style='List Bullet 2')

doc.add_heading('2.4 The Card List', level=2)
doc.add_paragraph(
    'Below the flashcard you will see a scrollable list of all loaded cards. '
    'Click on any card to jump directly to it. The speaker button next to each card reads it aloud individually. '
    'Cards with images are marked with a 🖼 icon in the list.'
)

# --- 3 ---
doc.add_heading('3. The Flexible Input Format', level=1)
doc.add_paragraph(
    'The app understands three input formats and automatically detects which one you are using. '
    'You can create vocabulary lists with any simple text editor (Notepad, TextEdit, etc.) and save them as .txt files.'
)

doc.add_heading('Format 1: Comma-separated (simplest format)', level=2)
doc.add_paragraph('The simplest format: Russian and English/German separated by a comma, one card per line.')
add_code_block(doc, 'дом, house\nшкола, school\nкнига, book\nкошка, cat\nсобака, dog')
doc.add_paragraph('Front = Russian (left of the comma), Back = translation (right of the comma).')

doc.add_heading('Format 2: Tab-separated (Quizlet-compatible)', level=2)
doc.add_paragraph(
    'When you export vocabulary lists from Quizlet, front and back are separated by a tab character. '
    'The app detects this automatically.'
)
add_code_block(doc, 'дом\thouse\nшкола\tschool\nкнига\tbook')
doc.add_paragraph('(Between the Russian and English word there is a tab character, i.e. the Tab key.)')

doc.add_heading('Format 3: Advanced format with additional information', level=2)

p = doc.add_paragraph()
r = p.add_run('Grammar hints')
r.bold = True
doc.add_paragraph('Place grammar rules in round brackets at the beginning of the line:')
add_code_block(doc, '(в + Accusative - direction: where to?)Я иду в школу., I am going to school.')
doc.add_paragraph('The grammar hint is displayed as a golden heading on the card.')

p = doc.add_paragraph()
r = p.add_run('Multi-line cards with example sentences')
r.bold = True
doc.add_paragraph('Use \\n in the text to create line breaks on the card:')
add_code_block(doc, '(в + Accusative)\\n"Я иду в школу.", I am going to school.')

p = doc.add_paragraph()
r = p.add_run('Controlling speech output')
r.bold = True
doc.add_paragraph('Put text in quotation marks to specify what exactly should be read aloud:')
add_code_block(doc, '1:00, "один час"')
doc.add_paragraph('Here 1:00 is displayed, but the speech synthesizer reads "один час".')

p = doc.add_paragraph()
r = p.add_run('Images on cards')
r.bold = True
doc.add_paragraph('Add images using the (img:path) tag:')
add_code_block(doc, '(img:images/dom.svg)дом, house')
doc.add_paragraph('The image is displayed on the front side above the Russian word. Supported formats: SVG, PNG, JPG, etc.')

doc.add_heading('Texts for Reading Aloud', level=2)
doc.add_paragraph('For the Read Texts mode, use plain running text. Paragraphs are separated by blank lines:')
add_code_block(doc, 'Здравствуйте. Вот, пожалуйста, ваше меню.\n\nЯ буду минеральную воду без газа.\n\nМне, пожалуйста, чёрный чай с лимоном.')

# --- 4 ---
doc.add_heading('4. Learning Modes — Spaced Repetition', level=1)

doc.add_heading('What is Spaced Repetition?', level=2)
doc.add_paragraph(
    'Spaced Repetition is a scientifically proven learning method: cards you know well are repeated less often, '
    'while cards you struggle with are shown more frequently. This optimizes your study time.'
)

doc.add_heading('The Three Modes at a Glance', level=2)
add_table(doc,
    ['Mode', 'Description', 'Recommended for'],
    [
        ['Free', 'Browse through all cards yourself', 'First introduction to new vocabulary'],
        ['SM-2', 'Classic system (since the 1980s)', 'Systematic long-term learning'],
        ['FSRS', 'Modern, optimized algorithm', 'Optimal long-term learning'],
    ]
)

doc.add_heading('How SM-2 / FSRS Works', level=2)
doc.add_paragraph('The app first shows you due cards (reviews), then new cards (max. 20 per session, adjustable).', style='List Number')
doc.add_paragraph('Look at the front side and try to recall the answer.', style='List Number')
doc.add_paragraph('Flip the card (click or Spacebar).', style='List Number')
doc.add_paragraph('Rate yourself honestly:', style='List Number')
doc.add_paragraph('Again — Shows the card again soon', style='List Bullet 2')
doc.add_paragraph('Hard — Short review interval', style='List Bullet 2')
doc.add_paragraph('Good — Normal interval (recommended for correct answers)', style='List Bullet 2')
doc.add_paragraph('Easy — Long interval (only for effortless recall)', style='List Bullet 2')
doc.add_paragraph('The app automatically calculates when you should see each card again.', style='List Number')

doc.add_heading('Statistics and Progress', level=2)
doc.add_paragraph(
    'Click on Statistics to see your learning progress: how many cards you have learned, '
    'how many reviews you have done, and your daily learning history. '
    'Progress is automatically saved in your browser.'
)

# --- 5 ---
doc.add_heading('5. Read Texts — Listening Comprehension', level=1)
doc.add_paragraph('Click on Read Texts and load a text file (Russian plain text).', style='List Number')
doc.add_paragraph('The text is displayed in a reading window.', style='List Number')
doc.add_paragraph('Choose the reading mode:', style='List Number')
doc.add_paragraph('All — The text is read sentence by sentence. The currently spoken word is highlighted in gold (word-by-word cursor). Already read sentences and words are dimmed.', style='List Bullet 2')
doc.add_paragraph('Sentence by sentence — Each sentence is read individually; click on a sentence to hear it. The current word is also highlighted. Use the ⏮ Previous / ⏭ Next buttons to navigate between sentences.', style='List Bullet 2')
doc.add_paragraph(
    'The text automatically scrolls to keep the current word visible at all times.'
)

# --- 6 ---
doc.add_heading('6. The Vocabulary Database', level=1)
doc.add_paragraph(
    'Every time you load a vocabulary file, the cards are automatically saved in a database in your browser. '
    'The database grows with each new file — without duplicates.'
)
doc.add_paragraph(
    'Click Export (Excel) to download your entire vocabulary collection as a TSV file. '
    'You can open this file directly in Excel or Google Sheets.'
)
doc.add_paragraph('The export file contains: Russian, German, Topic, Category, Grammar, and Source (filename).')

# --- 7 ---
doc.add_heading('7. Keyboard Shortcuts', level=1)
add_table(doc,
    ['Key', 'Function'],
    [
        ['Spacebar', 'Flip card'],
        ['Right arrow', 'Next card'],
        ['Left arrow', 'Previous card'],
        ['1', 'Rating: Again'],
        ['2', 'Rating: Hard'],
        ['3', 'Rating: Good'],
        ['4', 'Rating: Easy'],
    ]
)

# --- 8 ---
doc.add_heading('8. Using AI and LLMs to Create Vocabulary Lists', level=1)
doc.add_paragraph(
    'One of the greatest strengths of this app is its simple text format. You can ask ChatGPT, Claude, Gemini, '
    'or other AI assistants to create customized vocabulary lists for you — in seconds, tailored to your level and topic.'
)

# 8.1
doc.add_heading('8.1 Simple Vocabulary Lists', level=2)
p = doc.add_paragraph()
r = p.add_run('Prompt:')
r.bold = True
add_code_block(doc,
    'Create a Russian-English vocabulary list on the topic "Food and\n'
    'Drinks" for beginners (level A1). Format: one entry per line,\n'
    'Russian word (with stress marks), comma, English word.'
)
p = doc.add_paragraph()
r = p.add_run('Result (ready to use):')
r.bold = True
add_code_block(doc,
    'хлеб, bread\nмолоко́, milk\nмя́со, meat\nры́ба, fish\n'
    'о́вощи, vegetables\nфру́кты, fruit\nвода́, water\nчай, tea\n'
    'ко́фе, coffee\nсок, juice\nсыр, cheese\nма́сло, butter'
)

# 8.2
doc.add_heading('8.2 Vocabulary Lists with Grammar Hints', level=2)
p = doc.add_paragraph()
r = p.add_run('Prompt:')
r.bold = True
add_code_block(doc,
    'Create 10 flashcards about "Russian Verbs of Motion" with\n'
    'grammar hints. Format:\n\n'
    '(grammar hint)\\n"Russian example sentence", English translation\n\n'
    'Each card should have a grammar hint in brackets at the start,\n'
    'then \\n, then the Russian sentence in quotation marks,\n'
    'and after the comma the English translation. Use stress marks.'
)
p = doc.add_paragraph()
r = p.add_run('Result:')
r.bold = True
add_code_block(doc,
    '(идти́ - to go on foot, one-way)\\n"Я иду́ в магази́н.", I\'m going to the store.\n'
    '(ходи́ть - to go on foot, habitual)\\n"Я хожу́ в шко́лу ка́ждый день.", I go to school every day.\n'
    '(е́хать - to go by transport, one-way)\\n"Мы е́дем в Москву́.", We are going to Moscow.\n'
    '(е́здить - to go by transport, habitual)\\n"Он е́здит на рабо́ту на метро́.", He takes the metro to work.\n'
    '(лете́ть - to fly, one-way)\\n"Самолёт лети́т в Берли́н.", The plane is flying to Berlin.'
)

# 8.3
doc.add_heading('8.3 Extracting Vocabulary from a Text', level=2)
p = doc.add_paragraph()
r = p.add_run('Prompt:')
r.bold = True
add_code_block(doc,
    'Here is a Russian text from my textbook. Create a vocabulary list\n'
    'with all important words that an A1/A2 learner probably doesn\'t\n'
    'know yet. Format: Russian word with stress marks, comma, English.\n\n'
    'Text: "Здравствуйте. Вот, пожалуйста, ваше меню. Какие напитки\n'
    'желаете?"'
)
p = doc.add_paragraph()
r = p.add_run('Result:')
r.bold = True
add_code_block(doc,
    'меню́, menu\nнапи́тки, drinks\nжела́ть, to wish / to desire\n'
    'минера́льная вода́, mineral water\nбез газа́, still (without gas)\n'
    'чёрный чай, black tea\nлимо́н, lemon'
)

# 8.4
doc.add_heading('8.4 Quizlet Format (Tab-separated)', level=2)
p = doc.add_paragraph()
r = p.add_run('Prompt:')
r.bold = True
add_code_block(doc,
    'Create a Quizlet-compatible vocabulary list (tab-separated) with\n'
    '20 Russian adjectives and their English translations. Level A2.\n'
    'Format: Russian [TAB] English. Use stress marks.'
)
p = doc.add_paragraph()
r = p.add_run('Result (tab-separated):')
r.bold = True
add_code_block(doc,
    'большо́й\tbig\nма́ленький\tsmall\nно́вый\tnew\nста́рый\told\n'
    'молодо́й\tyoung\nкраси́вый\tbeautiful\nдо́брый\tkind\n'
    'плохо́й\tbad\nбы́стрый\tfast\nме́дленный\tslow'
)

# 8.5
doc.add_heading('8.5 Dialogue Cards with Full Sentences', level=2)
p = doc.add_paragraph()
r = p.add_run('Prompt:')
r.bold = True
add_code_block(doc,
    'Create 8 flashcards with typical sentences for a restaurant\n'
    'situation in Russian (level A2). Each card should have a Russian\n'
    'sentence on the front and the English translation on the back.\n'
    'Format: Russian, English — one entry per line. Use stress marks.'
)
p = doc.add_paragraph()
r = p.add_run('Result:')
r.bold = True
add_code_block(doc,
    'Мо́жно меню́, пожа́луйста?, Could I have the menu, please?\n'
    'Что вы рекоменду́ете?, What do you recommend?\n'
    'Я бу́ду стейк с карто́фелем., I\'ll have the steak with potatoes.\n'
    'Было́ о́чень вку́сно!, It was very delicious!\n'
    'Счёт, пожа́луйста., The bill, please.\n'
    'Мо́жно плати́ть ка́ртой?, Can I pay by card?'
)

# 8.6
doc.add_heading('8.6 Thematic Cards with Images', level=2)
p = doc.add_paragraph()
r = p.add_run('Prompt:')
r.bold = True
add_code_block(doc,
    'Create a vocabulary list on the topic "Family" with images.\n'
    'I have SVG images in the folder bilder_familie/ with files:\n'
    'mama.svg, papa.svg, syn.svg, doch.svg.\n\n'
    'Format: (img:bilder_familie/file.svg)Russian, English'
)
p = doc.add_paragraph()
r = p.add_run('Result:')
r.bold = True
add_code_block(doc,
    '(img:bilder_familie/mama.svg)ма́ма, mom\n'
    '(img:bilder_familie/papa.svg)па́па, dad\n'
    '(img:bilder_familie/syn.svg)сын, son\n'
    '(img:bilder_familie/doch.svg)дочь, daughter'
)

# 8.7
doc.add_heading('8.7 Creating a Text for Reading Aloud', level=2)
p = doc.add_paragraph()
r = p.add_run('Prompt:')
r.bold = True
add_code_block(doc,
    'Write a short Russian dialogue (about 10 sentences) on the topic\n'
    '"Shopping at the Supermarket" for level A2. The dialogue should be\n'
    'between a shop assistant and a customer. Russian text only, with\n'
    'paragraph breaks between speaker turns.'
)
p = doc.add_paragraph()
r = p.add_run('Result:')
r.bold = True
add_code_block(doc,
    'Продавщица: Здравствуйте! Могу я вам помочь?\n\n'
    'Покупатель: Да, пожалуйста. Где у вас молочные продукты?\n\n'
    'Продавщица: Молоко и сыр в третьем ряду, справа.\n\n'
    'Покупатель: Спасибо. А свежий хлеб у вас есть?\n\n'
    'Продавщица: Да, конечно. Хлеб вот здесь, рядом с кассой.'
)
doc.add_paragraph(
    'Save this text as a .txt file and load it in Read Texts mode — '
    'the app will read the dialogue sentence by sentence, with golden highlighting.'
)

# 8.8
doc.add_heading('8.8 Tips for Better AI Results', level=2)
doc.add_paragraph('Specify your level: Always mention your language level (A1, A2, B1, etc.).', style='List Bullet')
doc.add_paragraph('Request stress marks: Explicitly write "with stress marks" in your prompt.', style='List Bullet')
doc.add_paragraph('Describe the format precisely: Give the AI an example of what each line should look like.', style='List Bullet')
doc.add_paragraph('Narrow down the topic: The more specific, the more useful. Instead of "everyday life", try "daily routine of a student".', style='List Bullet')
doc.add_paragraph('Use your textbook: Copy a text into the prompt and ask the AI to extract vocabulary from it.', style='List Bullet')
doc.add_paragraph('Check for errors: Occasionally verify AI-generated content with a dictionary.', style='List Bullet')

# --- 9 ---
doc.add_heading('9. Tips and Notes', level=1)
doc.add_paragraph('Stress marks (e.g. молоко́) help with pronunciation and are displayed correctly.', style='List Bullet')
doc.add_paragraph('Available Russian voices depend on your browser. Chrome generally offers the best selection.', style='List Bullet')
doc.add_paragraph('All data stays local in your browser. No data is sent to any server.', style='List Bullet')
doc.add_paragraph('Use the mode switch buttons to switch directly between flashcard and text reader modes.', style='List Bullet')

# --- 10 ---
doc.add_heading('10. File Format Summary', level=1)
add_table(doc,
    ['Format', 'Example', 'When to use?'],
    [
        ['Comma-separated', 'дом, house', 'Simple vocabulary lists'],
        ['Tab-separated', 'дом [TAB] house', 'Quizlet export, large lists'],
        ['With grammar', '(Acc.)\\nsentence, translation', 'Grammar + example sentences'],
        ['With image', '(img:img.svg)word, translation', 'Visual learning'],
        ['Plain text', 'Paragraphs with blank lines', 'Read texts aloud'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Good luck learning Russian!')
r.bold = True
r.font.size = Pt(14)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

output_path = r'C:\Users\klaus\Documents\Russisch\Vokabeltrainer\User_Guide_EN.docx'
doc.save(output_path)
print(f'Word document saved: {output_path}')
