"""Erstellt die Gebrauchsanweisung als Word-Dokument."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# --- Styles anpassen ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)

# Hilfsfunktionen
def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()  # Abstand

# === INHALT ===

doc.add_heading('Russisch Vokabeltrainer — Gebrauchsanweisung', level=0)

doc.add_heading('Willkommen', level=1)
doc.add_paragraph(
    'Der Russisch Vokabeltrainer ist eine vielseitige Lern-App direkt im Browser. '
    'Sie brauchen nichts zu installieren — einfach die Datei index.html im Browser öffnen und loslegen.'
)
doc.add_paragraph('Die App bietet zwei Hauptfunktionen:')
doc.add_paragraph('Karteikarten — Vokabeln lernen mit virtuellen Lernkarten, die man umdrehen kann', style='List Bullet')
doc.add_paragraph('Texte vorlesen — Russische Texte werden Satz für Satz laut vorgelesen', style='List Bullet')
doc.add_paragraph(
    'Beide Funktionen nutzen die Sprachausgabe Ihres Browsers, um Ihnen die korrekte russische '
    'Aussprache vorzusprechen. So trainieren Sie gleichzeitig Lesen, Verstehen und Hören.'
)

# --- 1 ---
doc.add_heading('1. So starten Sie die App', level=1)
doc.add_paragraph('Öffnen Sie die Datei index.html in einem modernen Browser (Chrome, Edge oder Firefox).', style='List Number')
doc.add_paragraph(
    'Auf dem Startbildschirm sehen Sie:'
    , style='List Number')
doc.add_paragraph('Zwei große Schaltflächen: Karteikarten (das Kartensymbol) und Texte vorlesen (das Buchsymbol)', style='List Bullet 2')
doc.add_paragraph('Darunter die Stimmeinstellungen: Wählen Sie bereits hier Geschlecht, Stimme und Tempo für die Sprachausgabe — diese Einstellungen gelten für beide Modi.', style='List Bullet 2')
doc.add_paragraph('Oben rechts können Sie die Sprache der Benutzeroberfläche umschalten: DE (Deutsch), EN (Englisch) oder RU (Russisch).', style='List Number')

# --- 2 ---
doc.add_heading('2. Karteikarten — Vokabeln lernen', level=1)

doc.add_heading('2.1 Vokabeln laden', level=2)
doc.add_paragraph(
    'Klicken Sie auf Karteikarten. Sie haben nun zwei Möglichkeiten:'
)
p = doc.add_paragraph()
r = p.add_run('Neue Datei laden: ')
r.bold = True
p.add_run(
    'Laden Sie eine Textdatei mit Ihren Vokabeln. '
    'Sie können die Datei entweder per Drag & Drop in den Bereich ziehen oder durch Klick die Datei auswählen. '
    'Die App merkt sich den zuletzt verwendeten Ordner, sodass der Dateidialog beim nächsten Mal dort wieder öffnet (Chrome/Edge). '
    'Die App erkennt automatisch, welches Format Ihre Datei hat (siehe Abschnitt 3).'
)
p = doc.add_paragraph()
r = p.add_run('Gespeicherte Vokabeln laden: ')
r.bold = True
p.add_run(
    'Unterhalb des Datei-Upload-Bereichs sehen Sie die Liste Ihrer gespeicherten Vokabeln '
    'aus früheren Sitzungen (siehe Abschnitt 6). Sie können:'
)
doc.add_paragraph('Alle Vokabeln laden — Ihre gesamte Vokabelsammlung als Karteikarten', style='List Bullet')
doc.add_paragraph('Eine einzelne Lektion laden — klicken Sie auf den Titel (z.B. "Im Restaurant", "Antonyme"), '
                   'um nur die Vokabeln dieser Lektion zu lernen', style='List Bullet')

doc.add_heading('2.2 Einstellungen vor dem Start', level=2)
doc.add_paragraph('Nach dem Laden erscheint ein Einstellungsbildschirm:')
doc.add_paragraph('Lernmodus — Wählen Sie zwischen Frei, SM-2 oder FSRS (siehe Abschnitt 4)', style='List Bullet')
doc.add_paragraph('Auto-Vorlesen — Die Vorderseite wird automatisch vorgelesen', style='List Bullet')
doc.add_paragraph('Rückseite vorlesen — Auch die deutsche Übersetzung wird vorgelesen', style='List Bullet')
doc.add_paragraph('Klicken Sie auf Karteikarten starten, um zu beginnen.')

doc.add_heading('2.3 Während des Lernens', level=2)
doc.add_paragraph('Karte umdrehen: Klicken Sie auf die Karte oder drücken Sie die Leertaste.', style='List Bullet')
doc.add_paragraph('Nächste/Vorherige Karte: Benutzen Sie die Pfeil-Buttons oder die Pfeiltasten.', style='List Bullet')
doc.add_paragraph('Vorlesen: Klicken Sie auf den Lautsprecher-Button, um die aktuelle Karte vorlesen zu lassen.', style='List Bullet')
doc.add_paragraph('Bild hinzufügen: Klicken Sie auf den 📷-Button, um ein Bild von Ihrem Gerät zur aktuellen Karte hinzuzufügen. Das Bild erscheint auf der Vorderseite.', style='List Bullet')
doc.add_paragraph('Bild löschen: Fahren Sie mit der Maus über ein Bild auf der Karte — es erscheint ein kleiner ✕-Button, mit dem Sie das Bild wieder entfernen können.', style='List Bullet')
p = doc.add_paragraph(
    'Im Spaced-Repetition-Modus (SM-2 oder FSRS): Bewerten Sie nach dem Umdrehen, wie gut Sie die Karte wussten:'
)
doc.add_paragraph('Nochmal (Taste 1) — Wusste ich nicht', style='List Bullet 2')
doc.add_paragraph('Schwer (Taste 2) — Gerade so gewusst', style='List Bullet 2')
doc.add_paragraph('Gut (Taste 3) — Gewusst, aber musste nachdenken', style='List Bullet 2')
doc.add_paragraph('Leicht (Taste 4) — Sofort gewusst', style='List Bullet 2')

doc.add_heading('2.4 Die Kartenliste', level=2)
doc.add_paragraph(
    'Unterhalb der Karteikarte sehen Sie eine scrollbare Liste aller geladenen Karten. '
    'Klicken Sie auf eine Karte, um direkt dorthin zu springen. Der Lautsprecher-Button neben jeder Karte liest diese einzeln vor. '
    'Karten mit Bildern werden durch ein 🖼-Symbol in der Liste gekennzeichnet.'
)

# --- 3 ---
doc.add_heading('3. Das flexible Eingabeformat', level=1)
doc.add_paragraph(
    'Die App versteht drei Eingabeformate und erkennt automatisch, welches Sie verwenden. '
    'Sie können die Vokabellisten mit jedem einfachen Texteditor (Notepad, TextEdit usw.) erstellen und als .txt-Datei speichern.'
)

doc.add_heading('Format 1: Komma-getrennt (einfachstes Format)', level=2)
doc.add_paragraph('Das einfachste Format: Russisch und Deutsch durch ein Komma getrennt, eine Karte pro Zeile.')
add_code_block(doc, 'дом, Haus\nшкола, Schule\nкнига, Buch\nкошка, Katze\nсобака, Hund')
doc.add_paragraph('Vorderseite = Russisch (links vom Komma), Rückseite = Deutsch (rechts vom Komma).')

doc.add_heading('Format 2: Tab-getrennt (Quizlet-kompatibel)', level=2)
doc.add_paragraph(
    'Wenn Sie Vokabellisten von Quizlet exportieren, werden Vorder- und Rückseite durch einen Tabulator getrennt. '
    'Die App erkennt das automatisch.'
)
add_code_block(doc, 'дом\tHaus\nшкола\tSchule\nкнига\tBuch')
doc.add_paragraph('(Zwischen dem russischen und deutschen Wort steht ein Tabulator, also die Tab-Taste.)')

doc.add_heading('Format 3: Erweitertes Format mit Zusatzinformationen', level=2)

p = doc.add_paragraph()
r = p.add_run('Grammatikhinweise')
r.bold = True
doc.add_paragraph('Setzen Sie Grammatikregeln in runde Klammern am Anfang der Zeile:')
add_code_block(doc, '(в + Akkusativ - Richtung: wohin?)Я иду в школу., Ich gehe in die Schule.')
doc.add_paragraph('Der Grammatikhinweis wird als goldene Überschrift auf der Karte angezeigt.')

p = doc.add_paragraph()
r = p.add_run('Mehrzeilige Karten mit Beispielsätzen')
r.bold = True
doc.add_paragraph('Verwenden Sie \\n im Text, um Zeilenumbrüche auf der Karte zu erzeugen:')
add_code_block(doc, '(в + Akkusativ)\\n"Я иду в школу.", Ich gehe in die Schule.')

p = doc.add_paragraph()
r = p.add_run('Sprachausgabe steuern')
r.bold = True
doc.add_paragraph('Setzen Sie Text in Anführungszeichen, um festzulegen, was vorgelesen werden soll:')
add_code_block(doc, '1:00, "один час"')
doc.add_paragraph('Hier wird 1:00 angezeigt, aber die Sprachausgabe liest „один час" vor.')

p = doc.add_paragraph()
r = p.add_run('Bilder auf Karten')
r.bold = True
doc.add_paragraph('Fügen Sie Bilder hinzu mit dem Tag (img:Pfad):')
add_code_block(doc, '(img:bilder/dom.svg)дом, Haus')
doc.add_paragraph('Das Bild wird auf der Vorderseite über dem russischen Wort angezeigt. Unterstützt werden SVG, PNG, JPG usw.')

doc.add_heading('Texte zum Vorlesen', level=2)
doc.add_paragraph('Für den Modus Texte vorlesen verwenden Sie einfache Fließtexte. Absätze werden durch Leerzeilen getrennt:')
add_code_block(doc, 'Здравствуйте. Вот, пожалуйста, ваше меню.\n\nЯ буду минеральную воду без газа.\n\nМне, пожалуйста, чёрный чай с лимоном.')

# --- 4 ---
doc.add_heading('4. Lernmodi — Spaced Repetition', level=1)

doc.add_heading('Was ist Spaced Repetition?', level=2)
doc.add_paragraph(
    'Spaced Repetition ist eine wissenschaftlich erprobte Lernmethode: Karten, die Sie gut kennen, werden seltener wiederholt. '
    'Karten, die Ihnen schwerfallen, werden häufiger gezeigt. So nutzen Sie Ihre Lernzeit optimal.'
)

doc.add_heading('Die drei Modi im Überblick', level=2)
add_table(doc,
    ['Modus', 'Beschreibung', 'Empfohlen für'],
    [
        ['Frei', 'Selbst durch alle Karten blättern', 'Erstes Kennenlernen neuer Vokabeln'],
        ['SM-2', 'Klassisches System (seit den 1980ern)', 'Systematisches Langzeitlernen'],
        ['FSRS', 'Moderner, optimierter Algorithmus', 'Optimales Langzeitlernen'],
    ]
)

doc.add_heading('So funktioniert SM-2 / FSRS', level=2)
doc.add_paragraph('Die App zeigt Ihnen zuerst fällige Karten (Wiederholungen) und dann neue Karten (max. 20 pro Sitzung, einstellbar).', style='List Number')
doc.add_paragraph('Schauen Sie sich die Vorderseite an und versuchen Sie, die Antwort zu erinnern.', style='List Number')
doc.add_paragraph('Drehen Sie die Karte um (Klick oder Leertaste).', style='List Number')
doc.add_paragraph('Bewerten Sie sich ehrlich:', style='List Number')
doc.add_paragraph('Nochmal — Zeigt die Karte bald wieder', style='List Bullet 2')
doc.add_paragraph('Schwer — Kurzes Wiederholungsintervall', style='List Bullet 2')
doc.add_paragraph('Gut — Normales Intervall (empfohlen bei richtiger Antwort)', style='List Bullet 2')
doc.add_paragraph('Leicht — Langes Intervall (nur bei müheloser Antwort)', style='List Bullet 2')
doc.add_paragraph('Die App berechnet automatisch, wann Sie jede Karte wieder sehen sollten.', style='List Number')

doc.add_heading('Statistik und Fortschritt', level=2)
doc.add_paragraph(
    'Klicken Sie auf Statistik, um Ihren Lernfortschritt zu sehen: wie viele Karten Sie gelernt haben, '
    'wie viele Wiederholungen Sie durchgeführt haben, und Ihren täglichen Lernverlauf. '
    'Der Fortschritt wird automatisch im Browser gespeichert.'
)

# --- 5 ---
doc.add_heading('5. Texte vorlesen — Hörverstehen trainieren', level=1)
doc.add_paragraph('Klicken Sie auf Texte vorlesen und laden Sie eine Textdatei (Fließtext auf Russisch).', style='List Number')
doc.add_paragraph('Der Text wird in einem Lesefenster angezeigt.', style='List Number')
doc.add_paragraph('Wählen Sie den Lesemodus:', style='List Number')
doc.add_paragraph('ganzer Text — Der Text wird Satz für Satz vorgelesen. Das aktuell gesprochene Wort wird goldfarben markiert (Wort-für-Wort-Cursor). Bereits gelesene Sätze und Wörter werden ausgegraut.', style='List Bullet 2')
doc.add_paragraph('Satz für Satz — Jeder Satz wird einzeln vorgelesen; klicken Sie auf einen Satz, um ihn zu hören. Auch hier wird das aktuelle Wort markiert. Mit den Buttons ⏮ Voriger Satz / ⏭ Nächster Satz navigieren Sie zwischen den Sätzen.', style='List Bullet 2')
doc.add_paragraph(
    'Der Text scrollt automatisch mit, damit das aktuelle Wort immer im sichtbaren Bereich bleibt.'
)

# --- 6 ---
doc.add_heading('6. Die Vokabel-Datenbank — Ihr wachsender Wortschatz', level=1)

doc.add_heading('Automatisches Sammeln', level=2)
doc.add_paragraph(
    'Jedes Mal, wenn Sie eine Vokabeldatei laden oder Vokabeln per KI generieren lassen, werden die Karten automatisch in einer '
    'Datenbank im Browser gespeichert. Die Datenbank wächst mit jeder neuen Datei — ohne Duplikate. '
    'Jede Vokabel merkt sich, aus welcher Quelldatei (Lektion) sie stammt.'
)

doc.add_heading('Gespeicherte Vokabeln laden und lernen', level=2)
doc.add_paragraph(
    'Wenn Sie im Hauptmenü auf Karteikarten klicken, sehen Sie unterhalb des Datei-Upload-Bereichs Ihre gespeicherten Vokabeln:'
)
doc.add_paragraph('Alle Vokabeln (N Karten) — Lädt Ihre gesamte Vokabelsammlung als Karteikarten. Ideal für gemischtes Wiederholen.', style='List Bullet')
doc.add_paragraph('Einzelne Lektionen — Jede Datei, die Sie jemals geladen haben, erscheint als eigener Eintrag mit Kartenzähler. '
                   'Klicken Sie auf eine Lektion, um nur deren Vokabeln zu lernen.', style='List Bullet')
doc.add_paragraph(
    'So können Sie gezielt bestimmte Themen wiederholen, ohne die Originaldatei erneut laden zu müssen. '
    'Der Lernfortschritt (Spaced Repetition) bleibt dabei erhalten.'
)

doc.add_heading('Export für Excel', level=2)
doc.add_paragraph(
    'Klicken Sie im Einstellungsbildschirm auf Export (Excel), um Ihre gesamte Vokabelsammlung als CSV-Datei herunterzuladen. '
    'Diese Datei können Sie direkt in Excel oder Google Sheets öffnen.'
)
doc.add_paragraph('Die Exportdatei enthält: Russisch, Deutsch, Thema, Kategorie, Grammatik und Quelle (Dateiname).')

doc.add_heading('Auto-Export als JSON', level=2)
doc.add_paragraph(
    'In den Einstellungen (Zahnrad-Symbol oben rechts) können Sie einen Speicherort für eine JSON-Datei wählen. '
    'Die App exportiert dann bei jeder Änderung automatisch Ihre gesamte Vokabelsammlung in diese Datei. '
    'So haben Sie immer ein aktuelles Backup.'
)

# --- 7 ---
doc.add_heading('7. Tastenkürzel', level=1)
add_table(doc,
    ['Taste', 'Funktion'],
    [
        ['Leertaste', 'Karte umdrehen'],
        ['Pfeiltaste rechts', 'Nächste Karte'],
        ['Pfeiltaste links', 'Vorherige Karte'],
        ['1', 'Bewertung: Nochmal'],
        ['2', 'Bewertung: Schwer'],
        ['3', 'Bewertung: Gut'],
        ['4', 'Bewertung: Leicht'],
    ]
)

# --- 8 ---
doc.add_heading('8. KI und LLMs zur Erstellung von Vokabellisten', level=1)
doc.add_paragraph(
    'Eine der größten Stärken dieser App ist ihr einfaches Textformat. Sie können ChatGPT, Claude, Gemini oder andere '
    'KI-Assistenten bitten, Ihnen maßgeschneiderte Vokabellisten zu erstellen — in Sekunden, genau auf Ihr Niveau und Thema zugeschnitten.'
)

# 8.1
doc.add_heading('8.1 Einfache Vokabellisten erstellen', level=2)
p = doc.add_paragraph()
r = p.add_run('Prompt:')
r.bold = True
add_code_block(doc,
    'Erstelle mir eine Vokabelliste Russisch-Deutsch zum Thema\n'
    '"Essen und Trinken" für Anfänger (Niveau A1). Format: ein Eintrag\n'
    'pro Zeile, russisches Wort (mit Betonungszeichen), Komma,\n'
    'deutsches Wort.'
)
p = doc.add_paragraph()
r = p.add_run('Ergebnis (direkt nutzbar):')
r.bold = True
add_code_block(doc,
    'хлеб, Brot\nмолоко́, Milch\nмя́со, Fleisch\nры́ба, Fisch\n'
    'о́вощи, Gemüse\nфру́кты, Obst\nвода́, Wasser\nчай, Tee\n'
    'ко́фе, Kaffee\nсок, Saft\nсыр, Käse\nма́сло, Butter'
)

# 8.2
doc.add_heading('8.2 Vokabellisten mit Grammatikhinweisen', level=2)
p = doc.add_paragraph()
r = p.add_run('Prompt:')
r.bold = True
add_code_block(doc,
    'Erstelle mir 10 Karteikarten zum Thema "Russische Verben der\n'
    'Bewegung" mit Grammatikhinweisen. Format:\n\n'
    '(Grammatikhinweis)\\n"Russischer Beispielsatz", Deutsche Übersetzung\n\n'
    'Jede Karte soll einen Grammatikhinweis in Klammern am Anfang\n'
    'haben, dann \\n, dann den russischen Satz in Anführungszeichen\n'
    'und nach dem Komma die deutsche Übersetzung.'
)
p = doc.add_paragraph()
r = p.add_run('Ergebnis:')
r.bold = True
add_code_block(doc,
    '(идти́ - gehen, zu Fuß, einmalig)\\n"Я иду́ в магази́н.", Ich gehe in den Laden.\n'
    '(ходи́ть - gehen, regelmäßig)\\n"Я хожу́ в шко́лу ка́ждый день.", Ich gehe jeden Tag in die Schule.\n'
    '(е́хать - fahren, einmalig)\\n"Мы е́дем в Москву́.", Wir fahren nach Moskau.\n'
    '(е́здить - fahren, regelmäßig)\\n"Он е́здит на рабо́ту на метро́.", Er fährt mit der Metro zur Arbeit.\n'
    '(лете́ть - fliegen, einmalig)\\n"Самолёт лети́т в Берли́н.", Das Flugzeug fliegt nach Berlin.'
)

# 8.3
doc.add_heading('8.3 Vokabeln zu einem bestimmten Text erstellen', level=2)
p = doc.add_paragraph()
r = p.add_run('Prompt:')
r.bold = True
add_code_block(doc,
    'Hier ist ein russischer Text aus meinem Lehrbuch. Erstelle daraus\n'
    'eine Vokabelliste mit allen wichtigen Wörtern, die ein A1/A2-\n'
    'Lerner wahrscheinlich noch nicht kennt. Format: russisches Wort\n'
    'mit Betonungszeichen, Komma, deutsches Wort.\n\n'
    'Text: "Здравствуйте. Вот, пожалуйста, ваше меню. Какие напитки\n'
    'желаете?"'
)
p = doc.add_paragraph()
r = p.add_run('Ergebnis:')
r.bold = True
add_code_block(doc,
    'меню́, Speisekarte\nнапи́тки, Getränke\nжела́ть, wünschen\n'
    'минера́льная вода́, Mineralwasser\nбез газа́, ohne Kohlensäure\n'
    'чёрный чай, schwarzer Tee\nлимо́н, Zitrone'
)

# 8.4
doc.add_heading('8.4 Quizlet-Format (Tab-getrennt)', level=2)
p = doc.add_paragraph()
r = p.add_run('Prompt:')
r.bold = True
add_code_block(doc,
    'Erstelle eine Quizlet-kompatible Vokabelliste (Tab-getrennt) mit\n'
    '20 russischen Adjektiven und ihren deutschen Übersetzungen.\n'
    'Niveau A2. Format: Russisch [TAB] Deutsch.'
)
p = doc.add_paragraph()
r = p.add_run('Ergebnis (Tab-getrennt):')
r.bold = True
add_code_block(doc,
    'большо́й\tgroß\nма́ленький\tklein\nно́вый\tneu\nста́рый\talt\n'
    'молодо́й\tjung\nкраси́вый\tschön\nдо́брый\tgut / freundlich\n'
    'плохо́й\tschlecht\nбы́стрый\tschnell\nме́дленный\tlangsam'
)

# 8.5
doc.add_heading('8.5 Dialogkarten mit ganzen Sätzen', level=2)
p = doc.add_paragraph()
r = p.add_run('Prompt:')
r.bold = True
add_code_block(doc,
    'Erstelle 8 Karteikarten mit typischen Sätzen für eine Situation\n'
    'im Restaurant auf Russisch (Niveau A2). Jede Karte soll einen\n'
    'russischen Satz vorne und die deutsche Übersetzung hinten haben.\n'
    'Format: russisch, deutsch — ein Eintrag pro Zeile.'
)
p = doc.add_paragraph()
r = p.add_run('Ergebnis:')
r.bold = True
add_code_block(doc,
    'Мо́жно меню́, пожа́луйста?, Kann ich die Speisekarte haben, bitte?\n'
    'Что вы рекоменду́ете?, Was empfehlen Sie?\n'
    'Я бу́ду стейк с карто́фелем., Ich nehme das Steak mit Kartoffeln.\n'
    'Было́ о́чень вку́сно!, Es war sehr lecker!\n'
    'Счёт, пожа́луйста., Die Rechnung, bitte.\n'
    'Мо́жно плати́ть ка́ртой?, Kann ich mit Karte bezahlen?'
)

# 8.6
doc.add_heading('8.6 Thematische Karten mit Bildern', level=2)
p = doc.add_paragraph()
r = p.add_run('Prompt:')
r.bold = True
add_code_block(doc,
    'Erstelle eine Vokabelliste zum Thema "Familie" mit Bildern.\n'
    'Ich habe SVG-Bilder im Ordner bilder_familie/ mit folgenden\n'
    'Dateien: mama.svg, papa.svg, syn.svg, doch.svg.\n\n'
    'Format: (img:bilder_familie/datei.svg)russisch, deutsch'
)
p = doc.add_paragraph()
r = p.add_run('Ergebnis:')
r.bold = True
add_code_block(doc,
    '(img:bilder_familie/mama.svg)ма́ма, die Mama\n'
    '(img:bilder_familie/papa.svg)па́па, der Papa\n'
    '(img:bilder_familie/syn.svg)сын, der Sohn\n'
    '(img:bilder_familie/doch.svg)дочь, die Tochter'
)

# 8.7
doc.add_heading('8.7 Einen Vorlesetext erstellen lassen', level=2)
p = doc.add_paragraph()
r = p.add_run('Prompt:')
r.bold = True
add_code_block(doc,
    'Schreibe einen kurzen russischen Dialog (ca. 10 Sätze) zum Thema\n'
    '"Einkaufen im Supermarkt" für Niveau A2. Der Dialog soll zwischen\n'
    'einer Verkäuferin und einem Kunden stattfinden. Nur russischer\n'
    'Text, mit Absätzen zwischen den Sprecherwechseln.'
)
p = doc.add_paragraph()
r = p.add_run('Ergebnis:')
r.bold = True
add_code_block(doc,
    'Продавщица: Здравствуйте! Могу я вам помочь?\n\n'
    'Покупатель: Да, пожалуйста. Где у вас молочные продукты?\n\n'
    'Продавщица: Молоко и сыр в третьем ряду, справа.\n\n'
    'Покупатель: Спасибо. А свежий хлеб у вас есть?\n\n'
    'Продавщица: Да, конечно. Хлеб вот здесь, рядом с кассой.'
)
doc.add_paragraph(
    'Diesen Text speichern Sie als .txt-Datei und laden ihn im Modus Texte vorlesen — '
    'die App liest Ihnen den Dialog Satz für Satz vor, mit goldener Hervorhebung.'
)

# 8.8
doc.add_heading('8.8 Tipps für bessere KI-Ergebnisse', level=2)
doc.add_paragraph('Niveau angeben: Nennen Sie immer Ihr Sprachniveau (A1, A2, B1 usw.).', style='List Bullet')
doc.add_paragraph('Betonungszeichen verlangen: Schreiben Sie explizit „mit Betonungszeichen" in den Prompt.', style='List Bullet')
doc.add_paragraph('Format genau beschreiben: Geben Sie der KI ein Beispiel, wie eine Zeile aussehen soll.', style='List Bullet')
doc.add_paragraph('Thema eingrenzen: Je konkreter, desto nützlicher. Statt „Alltag" lieber „Tagesablauf eines Studenten".', style='List Bullet')
doc.add_paragraph('Direkt aus dem Lehrbuch: Kopieren Sie einen Text in den Prompt und bitten Sie die KI, Vokabeln zu extrahieren.', style='List Bullet')
doc.add_paragraph('Fehler prüfen: KI-generierte Inhalte gelegentlich mit einem Wörterbuch überprüfen.', style='List Bullet')

# --- 9 ---
doc.add_heading('9. Tipps und Hinweise', level=1)
doc.add_paragraph('Betonungszeichen (z.B. молоко́) helfen bei der Aussprache und werden korrekt angezeigt.', style='List Bullet')
doc.add_paragraph('Die verfügbaren russischen Stimmen hängen von Ihrem Browser ab. Chrome bietet die beste Auswahl.', style='List Bullet')
doc.add_paragraph('Alle Daten bleiben lokal in Ihrem Browser gespeichert. Es werden keine Daten an Server gesendet.', style='List Bullet')
doc.add_paragraph('Über die Modus-Wechsel-Buttons können Sie direkt zwischen Karteikarten und Text wechseln.', style='List Bullet')

# --- 10 ---
doc.add_heading('10. Zusammenfassung der Dateiformate', level=1)
add_table(doc,
    ['Format', 'Beispiel', 'Wann verwenden?'],
    [
        ['Komma-getrennt', 'дом, Haus', 'Einfache Vokabellisten'],
        ['Tab-getrennt', 'дом [TAB] Haus', 'Quizlet-Export, große Listen'],
        ['Mit Grammatik', '(Akk.)\\nSatz, Übersetzung', 'Grammatik + Beispielsätze'],
        ['Mit Bild', '(img:bild.svg)Wort, Übersetzung', 'Visuelles Lernen'],
        ['Fließtext', 'Absätze mit Leerzeilen', 'Texte vorlesen'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Viel Erfolg beim Russischlernen!')
r.bold = True
r.font.size = Pt(14)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Speichern
output_path = r'C:\Users\klaus\Desktop\Russisch\Vokabeltrainer\Gebrauchsanweisung.docx'
doc.save(output_path)
print(f'Word-Dokument gespeichert: {output_path}')
